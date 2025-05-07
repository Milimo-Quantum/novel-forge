from __future__ import annotations
import asyncio
from dataclasses import dataclass, field
import logging
import time
import os
import json
import re
from typing import List, Optional, Union, Annotated, Set

# Patch: ensure Edge is hashable for use in pydantic-graph type hints
@dataclass(frozen=True)
class Edge:
    label: str = ""

from pydantic_graph import BaseNode, End, GraphRunContext  

def extract_genre_from_text(text: str) -> str:
    """Extract genre string from agent output text."""
    try:
        # Look for JSON with 'genre' key
        match_json = re.search(r'\{.*\}', text, re.DOTALL)
        if match_json:
            import json as json_lib
            data = json_lib.loads(match_json.group(0))
            genre = data.get("genre") or data.get("Genre")
            if isinstance(genre, str) and genre.strip():
                return genre.strip()
        # Look for markdown or plain text patterns
        match = re.search(r"[Gg]enre[:：]\s*\**([^\n*]+)", text)
        if match:
            genre = match.group(1).strip()
            if genre:
                return genre
    except Exception:
        pass
    return "Unknown Genre"

def get_valid_genre(ctx) -> str:
    """Retrieve sanitized genre string, always lowercase, defaults to 'unknown genre'."""
    try:
        genre_raw = getattr(ctx.state, "genre", None)
        if isinstance(genre_raw, str) and genre_raw.strip():
            return genre_raw.strip().lower()
    except Exception:
        pass
    return "unknown genre"

# Ensure core pydantic_graph components are imported
from pydantic_graph import BaseNode, End, GraphRunContext  

# Assuming agents.py and tasks.py are structured appropriately
# We might need to adjust imports based on actual structure
from agents import BookAgents, BookResult # Assuming BookResult is still useful or adapted
from tasks import BookTasks
from graph_state import (
    BookGenerationState,
    ChapterResult,
    GenerationConfig,
    CharacterProfile,
    PolishingFeedback, # Import the new feedback model
    WritingStats # Import WritingStats for type checking
)
from utils import log_operation, logger # Assuming utils has this decorator

from opentelemetry import trace

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# === Enhancement Plan Nodes ===

from pydantic_graph import BaseNode, GraphRunContext  

from graph_state import BookGenerationState
from agents import BookAgents

from dataclasses import dataclass

@dataclass
class PlotScaffolding(BaseNode[BookGenerationState, BookAgents]):
    """Generates and updates a dynamic plot outline with key beats."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'PacingControl':
        ctx.state.update_progress("Generating Plot Scaffolding")
        tasks = BookTasks()
        try:
            outline_data = await tasks.create_book_outline(
                ctx.state.refined_concept,
                ctx.state.config.num_chapters,
                ctx.state.publishing_metadata
            )
            outline_text = outline_data.get("prompt", "") if isinstance(outline_data, dict) else str(outline_data)
            ctx.state.book_outline = outline_text
            ctx.state.publishing_metadata.plot_outline = outline_text
            logger.info("Plot scaffolding generated.")
        except Exception as e:
            ctx.state.publishing_metadata.plot_outline = f"Plot scaffolding failed: {e}"
            logger.error(f"Plot scaffolding failed: {e}")
        return PacingControl()

@dataclass
class PacingControl(BaseNode[BookGenerationState, BookAgents]):
    """Adjusts pacing metadata to enforce variation of intensity and tempo."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'SceneTransitionGenerator':
        ctx.state.update_progress("Adjusting Pacing")
        try:
            num_chapters = ctx.state.config.num_chapters
            genre_lower = get_valid_genre(ctx)
            genre_for_prompt = genre_lower.capitalize()
            if genre_lower == "unknown genre":
                logger.warning("Genre metadata missing or invalid, using 'unknown genre' for pacing plan.")
            synopsis = ctx.state.refined_concept or ctx.state.initial_concept or "No synopsis provided."
            preferences = getattr(ctx.state.publishing_metadata, "user_preferences", "No specific preferences provided.")

            # Compose prompt for LLM agent
            prompt = f"""You are an expert story pacing planner.

Given the following information, generate a detailed pacing plan for a novel:

GENRE:
{genre_for_prompt}

SYNOPSIS:
{synopsis}

USER PREFERENCES:
{preferences}

NUMBER OF CHAPTERS: {num_chapters}

For each chapter, assign a pacing label from this set (or similar nuanced terms):
- slow-burn
- world-building
- emotional
- medium
- high-intensity
- climax
- cooldown

Output the pacing plan as a JSON object mapping chapter numbers (starting from 1) to pacing labels.

Example:
{{
  "1": "slow-burn",
  "2": "world-building",
  "3": "medium",
  "4": "high-intensity",
  "5": "climax",
  "6": "cooldown"
}}

Only output the JSON object, with no extra commentary."""

            # Call LLM agent to generate pacing plan
            pacing_plan_str = await run_agent_task(
                ctx,
                "pacing_planner",
                lambda *_: {"prompt": prompt},
                update_status="Generating Dynamic Pacing Plan"
            )

            import json as json_lib
            pacing_labels = []
            try:
                pacing_dict = json_lib.loads(pacing_plan_str)
                if isinstance(pacing_dict, dict):
                    # Convert to list ordered by chapter number
                    pacing_labels = [pacing_dict.get(str(i), "medium") for i in range(1, num_chapters + 1)]
            except Exception:
                pacing_labels = []

            # Fallback if parsing failed or empty
            if not pacing_labels or len(pacing_labels) != num_chapters:
                logger.warning("LLM pacing plan invalid or incomplete, falling back to default pacing logic.")

                # Define act boundaries
                act1_end = max(1, num_chapters // 3)
                act2_end = max(act1_end + 1, 2 * num_chapters // 3)

                pacing_labels = []
                for ch_num in range(1, num_chapters + 1):
                    if ch_num <= act1_end:
                        act = 1
                    elif ch_num <= act2_end:
                        act = 2
                    else:
                        act = 3

                    if act == 1:
                        base_pacing = "quieter"
                    elif act == 2:
                        base_pacing = "medium"
                    else:
                        base_pacing = "high-intensity"

                    if "thriller" in genre_lower:
                        if act == 1:
                            base_pacing = "medium"
                        elif act == 2:
                            base_pacing = "high-intensity"
                        else:
                            base_pacing = "climax"
                    elif "romance" in genre_lower:
                        if act == 1:
                            base_pacing = "emotional"
                        elif act == 2:
                            base_pacing = "medium"
                        else:
                            base_pacing = "emotional-high"
                    elif "fantasy" in genre_lower:
                        if act == 1:
                            base_pacing = "world-building"
                        elif act == 2:
                            base_pacing = "medium"
                        else:
                            base_pacing = "high-intensity"

                    pacing_labels.append(base_pacing)

            ctx.state.publishing_metadata.pacing_labels = pacing_labels
            logger.info(f"Dynamic pacing labels set: {pacing_labels}")
        except Exception as e:
            logger.error(f"Pacing control failed: {e}")
        return SceneTransitionGenerator()

@dataclass
class SceneTransitionGenerator(BaseNode[BookGenerationState, BookAgents]):
    """Generates connective tissue paragraphs for smooth scene/chapter transitions."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'PlotArchitect':
        ctx.state.update_progress("Generating Scene Transitions")
        tasks = BookTasks()
        try:
            transitions = await tasks.enhance_narrative_flow(
                ctx.state.get_full_draft(),
                ctx.state.book_outline,
                ctx.state.config.num_chapters
            )
            ctx.state.publishing_metadata.scene_transitions = transitions
            logger.info("Scene transitions generated.")
        except Exception as e:
            ctx.state.publishing_metadata.scene_transitions = f"Scene transitions failed: {e}"
            logger.error(f"Scene transition generation failed: {e}")
        return PlotArchitect()

@dataclass
class MarketAnalysis(BaseNode[BookGenerationState, BookAgents]):
    """Analyzes the market landscape for the book genre."""
    async def run(
        self, ctx: GraphRunContext[BookGenerationState, BookAgents]
    ) -> Union[
        Annotated[GenrePositioning, Edge(label="Genre Positioning")],
        Annotated[AudienceTargeting, Edge(label="Audience Targeting")],
        Annotated[ComparativeTitles, Edge(label="Comparative Titles")]
    ]:
        trace.get_current_span().add_event("NodeStart", {"node": "MarketAnalysis"})
        freeze = ctx.state.freeze_flags.get("market_analysis", False)
        if freeze:
            logger.info("Market analysis frozen by user, skipping MarketAnalysis node.")
            ctx.state.update_progress("Market Analysis Frozen/Skipped")
            return GenrePositioning()

        ctx.state.update_progress("Market Analysis")
        tasks = BookTasks()
        try:
            prompt_data = await tasks.market_analysis(ctx.state.initial_concept)
            market_analysis = await run_agent_task(
                ctx,
                "market_analyst",
                lambda *_: prompt_data,
                update_status="Market Analysis"
            )
            print(f"[DEBUG] Raw market_analyst output: {market_analysis}")
            logger.warning(f"[DEBUG] Raw market_analyst output: {market_analysis}")
            ctx.state.publishing_metadata.market_analysis = market_analysis
            logger.info("Market analysis completed.")
        except Exception as e:
            ctx.state.publishing_metadata.market_analysis = f"Market analysis failed: {e}"
            logger.error(f"Market analysis failed: {e}")
        trace.get_current_span().add_event("NodeEnd", {"node": "MarketAnalysis"})
        return GenrePositioning()

@dataclass
class GenrePositioning(BaseNode[BookGenerationState, BookAgents]):
    """Determines optimal genre positioning."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> StrategicOutlineGenerator:
        trace.get_current_span().add_event("NodeStart", {"node": "GenrePositioning"})
        freeze = ctx.state.freeze_flags.get("genre_positioning", False)
        if freeze:
            logger.info("Genre positioning frozen by user, skipping GenrePositioning node.")
            ctx.state.update_progress("Genre Positioning Frozen/Skipped")
            return StrategicOutlineGenerator()

        ctx.state.update_progress("Genre Positioning")
        tasks = BookTasks()
        try:
            prompt_data = await tasks.genre_positioning(ctx.state.initial_concept)
            genre_positioning = await run_agent_task(
                ctx,
                "genre_specialist",
                lambda *_: prompt_data,
                update_status="Genre Positioning"
            )
            print(f"[DEBUG] Raw genre_specialist output: {genre_positioning}")
            logger.warning(f"[DEBUG] Raw genre_specialist output: {genre_positioning}")
            parsed_genre = extract_genre_from_text(genre_positioning)
            ctx.state.genre = parsed_genre
            ctx.state.publishing_metadata.genre_positioning = parsed_genre
            logger.info(f"Genre positioning completed. Extracted genre: {parsed_genre}")
        except Exception as e:
            ctx.state.publishing_metadata.genre_positioning = f"Genre positioning failed: {e}"
            logger.error(f"Genre positioning failed: {e}")
        trace.get_current_span().add_event("NodeEnd", {"node": "GenrePositioning"})
        return StrategicOutlineGenerator()

@dataclass
class AudienceTargeting(BaseNode[BookGenerationState, BookAgents]):
    """Identifies target audience segments."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> StrategicOutlineGenerator:
        trace.get_current_span().add_event("NodeStart", {"node": "AudienceTargeting"})
        freeze = ctx.state.freeze_flags.get("audience_targeting", False)
        if freeze:
            logger.info("Audience targeting frozen by user, skipping AudienceTargeting node.")
            ctx.state.update_progress("Audience Targeting Frozen/Skipped")
            return StrategicOutlineGenerator()

        ctx.state.update_progress("Audience Targeting")
        tasks = BookTasks()
        try:
            prompt_data = await tasks.audience_targeting(ctx.state.initial_concept)
            audience_profile = await run_agent_task(
                ctx,
                "audience_analyst",
                lambda *_: prompt_data,
                update_status="Audience Targeting"
            )
            ctx.state.publishing_metadata.audience = audience_profile
            logger.info("Audience targeting completed.")
        except Exception as e:
            ctx.state.publishing_metadata.audience = f"Audience targeting failed: {e}"
            logger.error(f"Audience targeting failed: {e}")
        trace.get_current_span().add_event("NodeEnd", {"node": "AudienceTargeting"})
        return StrategicOutlineGenerator()

@dataclass
class ComparativeTitles(BaseNode[BookGenerationState, BookAgents]):
    """Finds comparable titles for positioning."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> StrategicOutlineGenerator:
        trace.get_current_span().add_event("NodeStart", {"node": "ComparativeTitles"})
        freeze = ctx.state.freeze_flags.get("comparative_titles", False)
        if freeze:
            logger.info("Comparative titles frozen by user, skipping ComparativeTitles node.")
            ctx.state.update_progress("Comparative Titles Frozen/Skipped")
            return StrategicOutlineGenerator()

        ctx.state.update_progress("Comparative Titles")
        tasks = BookTasks()
        try:
            prompt_data = await tasks.comparative_titles(ctx.state.initial_concept)
            comps = await run_agent_task(
                ctx,
                "comparative_titles_researcher",
                lambda *_: prompt_data,
                update_status="Comparative Titles"
            )
            ctx.state.publishing_metadata.comparable_titles = comps
            logger.info("Comparative titles completed.")
        except Exception as e:
            ctx.state.publishing_metadata.comparable_titles = f"Comparative titles failed: {e}"
            logger.error(f"Comparative titles failed: {e}")
        trace.get_current_span().add_event("NodeEnd", {"node": "ComparativeTitles"})
        return StrategicOutlineGenerator()

@dataclass
class StrategicOutlineGenerator(BaseNode[BookGenerationState, BookAgents]):
    """Generates a strategic outline incorporating market insights."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'PlotArchitect':
        trace.get_current_span().add_event("NodeStart", {"node": "StrategicOutlineGenerator"})
        freeze = ctx.state.freeze_flags.get("strategic_outline", False)
        if freeze:
            logger.info("Strategic outline frozen by user, skipping StrategicOutlineGenerator node.")
            ctx.state.update_progress("Strategic Outline Frozen/Skipped")
            return PlotArchitect()

        ctx.state.update_progress("Strategic Outline Generation")
        tasks = BookTasks()
        try:
            prompt_data = await tasks.strategic_outline(
                ctx.state.initial_concept,
                getattr(ctx.state.publishing_metadata, "market_analysis", "")
            )
            strategic_outline = await run_agent_task(
                ctx,
                "strategic_outliner",
                lambda *_: prompt_data,
                update_status="Strategic Outline Generation"
            )
            ctx.state.publishing_metadata.strategic_outline = strategic_outline
            logger.info("Strategic outline generated.")
        except Exception as e:
            ctx.state.publishing_metadata.strategic_outline = f"Strategic outline failed: {e}"
            logger.error(f"Strategic outline generation failed: {e}")
        trace.get_current_span().add_event("NodeEnd", {"node": "StrategicOutlineGenerator"})
        return PlotArchitect()

@dataclass
class PlotArchitect(BaseNode[BookGenerationState, BookAgents]):
    """Designs the overall plot architecture and narrative arcs."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> CreateOutline:
        trace.get_current_span().add_event("NodeStart", {"node": "PlotArchitect"})
        ctx.state.update_progress("Plot Architecture")
        ctx.state.publishing_metadata.plot_architecture = "Plot architecture data."
        logger.info("Plot architecture completed.")
        result = CreateOutline()
        trace.get_current_span().add_event("NodeEnd", {"node": "PlotArchitect"})
        return result

@dataclass
class TransitionRefinement(BaseNode[BookGenerationState, BookAgents]):
    """Refines scene and chapter transitions for smooth narrative flow."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'CharacterVoiceRefinement':
        ctx.state.update_progress("Refining Transitions")
        agent = ctx.deps.transition_agent()
        content = ctx.state.get_full_draft()
        refined, _ = await agent.process_with_feedback(agent.run, content)
        ctx.state.final_book_content = refined
        return CharacterVoiceRefinement()

@dataclass
class CharacterVoiceRefinement(BaseNode[BookGenerationState, BookAgents]):
    """Refines character voices and dialogue realism."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'ThematicRefinement':
        ctx.state.update_progress("Refining Character Voices")
        agent = ctx.deps.character_voice_specialist()
        content = ctx.state.final_book_content or ctx.state.get_full_draft()
        refined, _ = await agent.process_with_feedback(agent.run, content)
        ctx.state.final_book_content = refined
        return ThematicRefinement()

@dataclass
class ThematicRefinement(BaseNode[BookGenerationState, BookAgents]):
    """Ensures varied yet coherent thematic expression."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'WorldBuildingRefinement':
        ctx.state.update_progress("Refining Themes")
        agent = ctx.deps.thematic_consistency_agent()
        content = ctx.state.final_book_content or ctx.state.get_full_draft()
        refined, _ = await agent.process_with_feedback(agent.run, content)
        ctx.state.final_book_content = refined
        return WorldBuildingRefinement()

@dataclass
class WorldBuildingRefinement(BaseNode[BookGenerationState, BookAgents]):
    """Injects concrete, sensory-rich world-building details."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> WriteCoordinator:
        ctx.state.update_progress("Refining World-Building")
        agent = ctx.deps.world_building_agent()
        content = ctx.state.final_book_content or ctx.state.get_full_draft()
        refined, _ = await agent.process_with_feedback(agent.run, content)
        ctx.state.final_book_content = refined
        return WriteCoordinator()

@dataclass
class WriteCoordinator(BaseNode[BookGenerationState, BookAgents]):
    """Coordinates multiple writer agents and manages writing tasks."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> WriteChapter:
        trace.get_current_span().add_event("NodeStart", {"node": "WriteCoordinator"})
        ctx.state.update_progress("Coordinating Writing")
        logger.info("Starting enhanced multi-agent write coordination...")

        tasks = BookTasks()

        outline = ctx.state.book_outline
        characters = ctx.state.characters
        world = ctx.state.world_details
        num_chapters = ctx.state.config.num_chapters
        min_words = ctx.state.config.min_words_per_chapter
        state_obj = ctx.state

        # --- Initial parallel generation ---
        async def run_writer():
            prompt_data = TaskPrompts.write_chapter_prompt(ctx.state, 1, min_words)
            return await run_agent_task(
                ctx, "writer", lambda *_: prompt_data,
                update_status="Coordinating Writing (Base Narrative)"
            )

        async def run_description():
            prompt_data = TaskPrompts.write_chapter_prompt(ctx.state, 1, min_words)
            return await run_agent_task(
                ctx, "description_architect", lambda *_: prompt_data,
                update_status="Coordinating Writing (Descriptions)"
            )

        async def run_dialogue():
            prompt_data = TaskPrompts.write_chapter_prompt(ctx.state, 1, min_words)
            return await run_agent_task(
                ctx, "dialogue_specialist", lambda *_: prompt_data,
                update_status="Coordinating Writing (Dialogue)"
            )

        base_narrative, description_content, dialogue_content = await asyncio.gather(
            run_writer(), run_description(), run_dialogue()
        )

        ctx.state.user_feedback["initial_base_narrative"] = base_narrative
        ctx.state.user_feedback["initial_description"] = description_content
        ctx.state.user_feedback["initial_dialogue"] = dialogue_content

        # --- Enhanced iterative explicit feedback refinement ---
        writer_agent = ctx.deps.writer()
        description_agent = ctx.deps.description_architect()
        dialogue_agent = ctx.deps.dialogue_specialist()

        refined_base = base_narrative
        refined_description = description_content
        refined_dialogue = dialogue_content

        max_iterations = getattr(ctx.state.config, 'max_iterations', 2)
        all_structured_feedback = []

        prev_base = ""
        prev_desc = ""
        prev_dial = ""

        for iteration in range(1, max_iterations + 1):
            feedback_focus = [
                {"category": "transition clarity", "priority": "medium", "suggestion": "Improve scene transitions using sensory anchors."},
                {"category": "character motivation", "priority": "high", "suggestion": "Clarify character desires, fears, and relationships in actions/dialogue."},
                {"category": "emotional resonance", "priority": "high", "suggestion": "Add vulnerability moments and symbolic objects."},
                {"category": "world integration", "priority": "medium", "suggestion": "Embed tactile details and respect world rules."},
                {"category": "thematic depth", "priority": "medium", "suggestion": "Tie themes to character choices and stakes."},
                {"category": "pacing", "priority": "medium", "suggestion": "Follow beat sheet: Goal → Obstacle → Choice → Consequence."}
            ]

            # Cap and filter feedback list: Only keep last 2 rounds, prioritize high/critical
            combined_feedback = all_structured_feedback[-6:] + feedback_focus
            combined_feedback = [
                x.model_dump() if hasattr(x, "model_dump") else x
                for x in combined_feedback
            ]
            combined_feedback = sorted(combined_feedback, key=lambda x: x.get("priority", ""), reverse=True)
            combined_feedback = combined_feedback[:6]

            # Writer refines
            refined_base, writer_feedback = await writer_agent.process_with_feedback(
                writer_agent.run,
                refined_base,
                feedback_data={
                    "description_feedback": refined_description,
                    "dialogue_feedback": refined_dialogue
                },
                feedback_list=combined_feedback
            )
            all_structured_feedback.extend(writer_feedback)

            # Description refines
            refined_description, description_feedback = await description_agent.process_with_feedback(
                description_agent.run,
                refined_description,
                feedback_data={
                    "writer_feedback": refined_base,
                    "dialogue_feedback": refined_dialogue
                },
                feedback_list=combined_feedback
            )
            all_structured_feedback.extend(description_feedback)

            # Dialogue refines
            refined_dialogue, dialogue_feedback = await dialogue_agent.process_with_feedback(
                dialogue_agent.run,
                refined_dialogue,
                feedback_data={
                    "writer_feedback": refined_base,
                    "description_feedback": refined_description
                },
                feedback_list=combined_feedback
            )
            all_structured_feedback.extend(dialogue_feedback)

            ctx.state.user_feedback[f"refined_base_round_{iteration}"] = refined_base
            ctx.state.user_feedback[f"refined_description_round_{iteration}"] = refined_description
            ctx.state.user_feedback[f"refined_dialogue_round_{iteration}"] = refined_dialogue

            # Compute normalized content change ratios
            def diff_ratio(a, b):
                a_tokens = set(a.split())
                b_tokens = set(b.split())
                if not a_tokens and not b_tokens:
                    return 0.0
                return len(a_tokens.symmetric_difference(b_tokens)) / max(len(a_tokens), 1)

            base_change = diff_ratio(prev_base, refined_base)
            desc_change = diff_ratio(prev_desc, refined_description)
            dial_change = diff_ratio(prev_dial, refined_dialogue)

            logger.info(f"Iteration {iteration} content change ratios - Base: {base_change:.3f}, Desc: {desc_change:.3f}, Dial: {dial_change:.3f}")
            logger.info(f"Feedback list size: {len(combined_feedback)}. Context size (tokens): Base: {len(refined_base.split())}, Desc: {len(refined_description.split())}, Dial: {len(refined_dialogue.split())}")
            if len(combined_feedback) > 8:
                logger.warning("Feedback list exceeds recommended size; context window may be at risk.")
            if max(len(refined_base.split()), len(refined_description.split()), len(refined_dialogue.split())) > 2048:
                logger.warning("Agent context size exceeds 2048 tokens; risk of LLM context loss.")

            # Early stop if all changes are minimal
            if base_change < 0.02 and desc_change < 0.02 and dial_change < 0.02:
                logger.info(f"Convergence detected at iteration {iteration}, stopping early.")
                break

            # Update previous outputs
            prev_base = refined_base
            prev_desc = refined_description
            prev_dial = refined_dialogue

            logger.info(f"Completed WriteCoordinator refinement iteration {iteration}")

        # After iterations, store accumulated structured feedback (last 6 only)
        if hasattr(ctx.state, "feedback_history") and isinstance(ctx.state.feedback_history, list):
            ctx.state.feedback_history.extend(all_structured_feedback[-6:])
        else:
            ctx.state.feedback_history = all_structured_feedback[-6:]

        # Merge final refined outputs
        composite_draft = f"{refined_base}\n\n---\n\n{refined_description}\n\n---\n\n{refined_dialogue}"
        ctx.state.user_feedback["coordinated_draft"] = composite_draft

        logger.info("Enhanced multi-agent write coordination with explicit iterative feedback completed.")
        ctx.state.update_progress("Multi-Agent Draft Created")

        result = WriteChapter(chapter_num=1)
        trace.get_current_span().add_event("NodeEnd", {"node": "WriteCoordinator"})
        return result

@dataclass
class MultiStageReviewHub(BaseNode[BookGenerationState, BookAgents]):
    """Coordinates multiple parallel review processes."""
    async def run(
        self, ctx: GraphRunContext[BookGenerationState, BookAgents]
    ) -> Union[
        Annotated[PeerReview, Edge(label="Peer Review")],
        Annotated[EditorialReview, Edge(label="Editorial Review")],
        Annotated[ConsistencyCheck, Edge(label="Consistency Check")],
        Annotated[StyleRefinement, Edge(label="Style Refinement")],
        Annotated[FlowEnhancement, Edge(label="Flow Enhancement")]
    ]:
        trace.get_current_span().add_event("NodeStart", {"node": "MultiStageReviewHub"})
        ctx.state.update_progress("Multi-Stage Review Hub")

        # Optionally, could still run all in parallel here, but explicit graph edges model parallelism
        return PeerReview()

@dataclass
class StructuralEditor(BaseNode[BookGenerationState, BookAgents]):
    """Performs structural editing of the manuscript."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> LineEditor:
        trace.get_current_span().add_event("NodeStart", {"node": "StructuralEditor"})
        ctx.state.update_progress("Structural Editing")
        manuscript = ctx.state.final_book_content or ctx.state.book_outline or ""
        prompt_data = TaskPrompts.structural_editing(manuscript)
        _ = await run_agent_task(
            ctx, "structural_editor", lambda *_: prompt_data,
            update_status="Structural Editing"
        )
        logger.info("Structural editing completed.")
        result = LineEditor()
        trace.get_current_span().add_event("NodeEnd", {"node": "StructuralEditor"})
        return result

@dataclass
class LineEditor(BaseNode[BookGenerationState, BookAgents]):
    """Performs line editing of the manuscript."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> PeerReviewSimulator:
        trace.get_current_span().add_event("NodeStart", {"node": "LineEditor"})
        ctx.state.update_progress("Line Editing")
        manuscript = ctx.state.final_book_content or ctx.state.book_outline or ""
        prompt_data = TaskPrompts.line_editing(manuscript)
        _ = await run_agent_task(
            ctx, "line_editor", lambda *_: prompt_data,
            update_status="Line Editing"
        )
        logger.info("Line editing completed.")
        result = PeerReviewSimulator()
        trace.get_current_span().add_event("NodeEnd", {"node": "LineEditor"})
        return result

@dataclass
class PeerReviewSimulator(BaseNode[BookGenerationState, BookAgents]):
    """Simulates peer review feedback."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> MultiStageReviewHub:
        ctx.state.update_progress("Peer Review Simulation")
        manuscript = ctx.state.final_book_content or ctx.state.book_outline or ""
        prompt_data = TaskPrompts.peer_review_simulation(manuscript)
        _ = await run_agent_task(
            ctx, "peer_review_simulator", lambda *_: prompt_data,
            update_status="Peer Review Simulation"
        )
        logger.info("Peer review simulation completed.")
        return MultiStageReviewHub()

@dataclass
class StyleGuideEnforcer(BaseNode[BookGenerationState, BookAgents]):
    """Checks and enforces style guide compliance."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> FormattingOptimizer:
        ctx.state.update_progress("Style Guide Enforcement")
        logger.info("Style guide enforcement completed.")
        return FormattingOptimizer()

@dataclass
class FormattingOptimizer(BaseNode[BookGenerationState, BookAgents]):
    """Optimizes formatting for publishing standards."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> MarketMetadataGenerator:
        ctx.state.update_progress("Formatting Optimization")
        logger.info("Formatting optimization completed.")
        return MarketMetadataGenerator()

@dataclass
class MarketMetadataGenerator(BaseNode[BookGenerationState, BookAgents]):
    """Generates metadata for publishing platforms."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> GenerateFrontMatter:
        ctx.state.update_progress("Market Metadata Generation")

        concept = ctx.state.refined_concept or ctx.state.initial_concept
        outline = ctx.state.book_outline or ""
        prompt_data = TaskPrompts.metadata_generation(concept)
        metadata = await run_agent_task(
            ctx, "metadata_generator", lambda *_: prompt_data,
            update_status="Market Metadata Generation"
        )
        ctx.state.publishing_metadata.metadata = metadata
        logger.info("Market metadata generated and stored.")
        return GenerateFrontMatter()

@dataclass
class GenerateProposal(BaseNode[BookGenerationState, BookAgents]):
    """Generates a publishing proposal with market analysis, audience targeting, comparable titles."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> GenerateTitle:
        ctx.state.update_progress("Generating Publishing Proposal")
        tasks = BookTasks()
        try:
            # Compose a detailed prompt integrating all prior analyses
            market_data = getattr(ctx.state.publishing_metadata, "market_analysis", "")
            genre_data = getattr(ctx.state.publishing_metadata, "genre_positioning", "")
            audience_data = getattr(ctx.state.publishing_metadata, "audience", "")
            comps_data = getattr(ctx.state.publishing_metadata, "comparable_titles", "")
            strategic_outline = getattr(ctx.state.publishing_metadata, "strategic_outline", "")

            prompt_data = {
                "prompt": f"""Based on the following analyses, generate a detailed publishing proposal:

MARKET ANALYSIS:
{market_data}

GENRE POSITIONING:
{genre_data}

AUDIENCE PROFILE:
{audience_data}

COMPARABLE TITLES:
{comps_data}

STRATEGIC OUTLINE:
{strategic_outline}

Include:
- Unique selling points
- Market fit
- Audience appeal
- Competitive differentiation
- Suggested positioning and marketing angles
- Potential challenges and mitigation strategies

Return a comprehensive proposal in markdown format."""
            }

            proposal = await run_agent_task(
                ctx,
                "market_analyst",
                lambda *_: prompt_data,
                update_status="Generating Publishing Proposal"
            )
            ctx.state.publishing_metadata.proposal = proposal
            logger.info("Generated publishing proposal metadata.")
        except Exception as e:
            ctx.state.publishing_metadata.proposal = f"Proposal generation failed: {e}"
            logger.error(f"Proposal generation failed: {e}")
        return GenerateTitle()

@dataclass
class GenerateSummaries(BaseNode[BookGenerationState, BookAgents]):
    """Generates summaries of chapters and the overall book to support context management."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> StructuralEditor:
        try:
            ctx.state.update_progress("Generating Summaries")
            summaries = {}
            combined_summary_parts = []
            for ch_num in range(1, ctx.state.config.num_chapters + 1):
                chapter = ctx.state.get_latest_chapter_version(ch_num)
                if not chapter:
                    continue
                # Simple heuristic: first 3 sentences as summary
                sentences = re.split(r'(?<=[.!?]) +', chapter.content)
                summary = ' '.join(sentences[:3]).strip()
                summaries[ch_num] = summary
                combined_summary_parts.append(f"Chapter {ch_num}: {summary}")
            ctx.state.chapter_summaries = summaries
            # Sanitize: replace None with empty string
            ctx.state.chapter_summaries = {
                k: (v if isinstance(v, str) and v is not None else "")
                for k, v in ctx.state.chapter_summaries.items()
            }
            ctx.state.combined_summary = "\n".join(combined_summary_parts)
            logger.info("Generated chapter and combined summaries.")
            return StructuralEditor()
        except Exception as e:
            logger.error(f"Error generating summaries: {e}")
            ctx.state.error_message = f"Error generating summaries: {e}"
            raise

@dataclass
class GenerateFrontMatter(BaseNode[BookGenerationState, BookAgents]):
    """Generates front matter content like copyright, dedication, about author."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> GenerateBackMatter:
        try:
            ctx.state.update_progress("Generating Front Matter")
            # Placeholder: generate content based on metadata
            meta = ctx.state.publishing_metadata or {}
            ctx.state.copyright_page = f"  {getattr(ctx.state.publishing_metadata, 'author', 'Author')} {time.strftime('%Y')}. All rights reserved."
            ctx.state.dedication_page = getattr(ctx.state.publishing_metadata, 'dedication', None) or "To all readers and dreamers."
            ctx.state.about_author_page = getattr(ctx.state.publishing_metadata, 'about_author', None) or f"{getattr(ctx.state.publishing_metadata, 'author', 'The author')} is a passionate storyteller."
            ctx.state.front_matter = f"{ctx.state.copyright_page}\n\n{ctx.state.dedication_page}\n\n{ctx.state.about_author_page}"
            logger.info("Generated front matter content.")
            return GenerateBackMatter()
        except Exception as e:
            logger.error(f"Error generating front matter: {e}")
            ctx.state.error_message = f"Error generating front matter: {e}"
            raise

@dataclass
class GenerateBackMatter(BaseNode[BookGenerationState, BookAgents]):
    """Generates back matter content like acknowledgments."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> AssembleBook:
        try:
            ctx.state.update_progress("Generating Back Matter")
            # Placeholder: generate acknowledgments
            ctx.state.acknowledgments_page = "Special thanks to everyone who supported this book's creation."
            ctx.state.back_matter = ctx.state.acknowledgments_page
            logger.info("Generated back matter content.")
            return AssembleBook()
        except Exception as e:
            logger.error(f"Error generating back matter: {e}")
            ctx.state.error_message = f"Error generating back matter: {e}"
            raise

@dataclass
class EvaluateQualityMetrics(BaseNode[BookGenerationState, BookAgents]):
    """Evaluates advanced quality metrics and readability of the manuscript."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> StyleGuideEnforcer:
        try:
            content = ctx.state.get_full_draft()
            metrics_report = ""

            try:
                import textstat
                fk_grade = textstat.flesch_kincaid_grade(content)
                fk_ease = textstat.flesch_reading_ease(content)
                smog = textstat.smog_index(content)
                coleman = textstat.coleman_liau_index(content)
                metrics_report += f"Flesch-Kincaid Grade Level: {fk_grade:.2f}\n"
                metrics_report += f"Flesch Reading Ease: {fk_ease:.2f}\n"
                metrics_report += f"SMOG Index: {smog:.2f}\n"
                metrics_report += f"Coleman-Liau Index: {coleman:.2f}\n"
            except ImportError:
                metrics_report += "textstat not installed, skipping readability scores.\n"
                fk_grade = fk_ease = smog = coleman = 0
            except Exception as e:
                metrics_report += f"Error calculating readability: {e}\n"
                fk_grade = fk_ease = smog = coleman = 0

            # Genre-specific heuristics (placeholder logic)
            genre = getattr(ctx.state.publishing_metadata, "genre", None) if ctx.state.publishing_metadata else None
            if genre:
                metrics_report += f"Genre: {genre}\n"
            else:
                metrics_report += "Genre: Unknown\n"

            # Heuristic: dialogue ratio
            dialogue_lines = len([line for line in content.splitlines() if line.strip().startswith('"') or line.strip().startswith("'")])
            total_lines = len(content.splitlines())
            dialogue_ratio = (dialogue_lines / total_lines) * 100 if total_lines else 0
            metrics_report += f"Dialogue Ratio: {dialogue_ratio:.1f}%\n"

            # Heuristic: description density (number of adjectives)
            import re as regex
            adjectives = regex.findall(r'\b\w+ly\b|\b\w+ous\b|\b\w+ive\b|\b\w+al\b', content)
            adj_density = (len(adjectives) / max(1, len(content.split()))) * 100
            metrics_report += f"Description Density (adj/words): {adj_density:.2f}%\n"

            # Placeholder: engagement score, style consistency
            try:
                # Example heuristic: ratio of exclamation marks as engagement proxy
                exclamations = content.count("!")
                engagement_score = min(10, exclamations / max(1, len(content) / 1000))
                metrics_report += f"Engagement Score (heuristic): {engagement_score:.2f}\n"
            except Exception:
                engagement_score = 0
                metrics_report += "Engagement Score: Error calculating\n"

            try:
                # Example heuristic: ratio of varied sentence lengths as style consistency proxy
                sentences = re.split(r'[.!?]', content)
                lengths = [len(s.split()) for s in sentences if s.strip()]
                if lengths:
                    avg_len = sum(lengths) / len(lengths)
                    variance = sum((l - avg_len) ** 2 for l in lengths) / len(lengths)
                    style_consistency = max(0, 10 - variance / 10)
                else:
                    style_consistency = 0
                metrics_report += f"Style Consistency (heuristic): {style_consistency:.2f}\n"
            except Exception:
                style_consistency = 0
                metrics_report += "Style Consistency: Error calculating\n"

            ctx.state.error_message = None  # Clear previous errors
            ctx.state.update_progress("Quality Metrics Evaluation Complete")
            logger.info(f"Quality Metrics:\n{metrics_report}")

            # --- Integration: Trigger rewrites based on metrics thresholds ---
            thresholds = {
                "fk_grade_min": 4.0,  # Minimum grade level
                "fk_ease_min": 50.0,  # Minimum reading ease
                "dialogue_ratio_min": 5.0,  # At least 5% dialogue
                "adj_density_min": 1.0,  # At least 1% adjectives
                "engagement_min": 1.0,  # Placeholder
                "style_consistency_min": 1.0  # Placeholder
            }

            needs_rewrite = False
            reasons = []

            if fk_grade < thresholds["fk_grade_min"]:
                needs_rewrite = True
                reasons.append(f"Flesch-Kincaid grade too low ({fk_grade:.2f})")
            if fk_ease < thresholds["fk_ease_min"]:
                needs_rewrite = True
                reasons.append(f"Reading ease too low ({fk_ease:.2f})")
            if dialogue_ratio < thresholds["dialogue_ratio_min"]:
                needs_rewrite = True
                reasons.append(f"Dialogue ratio too low ({dialogue_ratio:.1f}%)")
            if adj_density < thresholds["adj_density_min"]:
                needs_rewrite = True
                reasons.append(f"Description density too low ({adj_density:.2f}%)")
            if engagement_score < thresholds["engagement_min"]:
                needs_rewrite = True
                reasons.append("Engagement score too low (placeholder)")
            if style_consistency < thresholds["style_consistency_min"]:
                needs_rewrite = True
                reasons.append("Style consistency too low (placeholder)")

            if needs_rewrite:
                ctx.state.review_feedback = "Quality metrics below thresholds:\n" + "\n".join(reasons)
                # Mark all chapters as needing rewrite (or refine to specific chapters later)
                for ch_num in range(1, ctx.state.config.num_chapters + 1):
                    ctx.state.current_chapter_status[ch_num] = "needs_rewrite"
                logger.info("Marked chapters for rewrite due to low quality metrics.")
            else:
                logger.info("Quality metrics acceptable, no rewrite triggered.")

            return StyleGuideEnforcer()
        except Exception as e:
            logger.error(f"Error evaluating quality metrics: {e}")
            ctx.state.error_message = f"Error evaluating quality metrics: {e}"
            raise

@dataclass
class VerifyCrossReferences(BaseNode[BookGenerationState, BookAgents]):
    """Verifies, annotates, and auto-updates cross-references between chapters, characters, and world elements."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> GenerateSummaries:
        try:
            inconsistencies = []
            crossref_map = {}
            for ch_num, versions in ctx.state.chapter_versions.items():
                latest = ctx.state.get_latest_chapter_version(ch_num)
                if not latest:
                    continue
                refs = latest.references or {}
                updated_refs = {}
                annotated_content = latest.content

                # Auto-update references based on current characters/world
                for ref_key, ref_desc in refs.items():
                    found = False
                    # Fuzzy match characters
                    for char in ctx.state.characters:
                        if ref_key.lower() in char.name.lower() or char.name.lower() in ref_key.lower():
                            updated_refs[char.name] = char.description
                            found = True
                            annotated_content = re.sub(
                                rf"\b{re.escape(ref_key)}\b",
                                f"[{char.name}](#character-{char.name.replace(' ', '-')})",
                                annotated_content
                            )
                            crossref_map.setdefault(ch_num, []).append({"ref": char.name, "type": "character"})
                            break
                    # Fuzzy match world details
                    if not found and ctx.state.world_details:
                        if ref_key.lower() in ctx.state.world_details.lower():
                            updated_refs[ref_key] = "World element"
                            found = True
                            annotated_content = re.sub(
                                rf"\b{re.escape(ref_key)}\b",
                                f"[{ref_key}](#world-{ref_key.replace(' ', '-')})",
                                annotated_content
                            )
                            crossref_map.setdefault(ch_num, []).append({"ref": ref_key, "type": "world"})
                    if not found:
                        inconsistencies.append(f"Chapter {ch_num}: Reference '{ref_key}' not found in characters or world.")

                # Auto-add missing references from current characters/world
                for char in ctx.state.characters:
                    if char.name not in updated_refs:
                        if char.name in annotated_content:
                            updated_refs[char.name] = char.description
                            annotated_content = re.sub(
                                rf"\b{re.escape(char.name)}\b",
                                f"[{char.name}](#character-{char.name.replace(' ', '-')})",
                                annotated_content
                            )
                            crossref_map.setdefault(ch_num, []).append({"ref": char.name, "type": "character"})
                if ctx.state.world_details:
                    world_terms = re.findall(r'\b\w+\b', ctx.state.world_details)
                    for term in world_terms:
                        if term not in updated_refs and term in annotated_content:
                            updated_refs[term] = "World element"
                            annotated_content = re.sub(
                                rf"\b{re.escape(term)}\b",
                                f"[{term}](#world-{term.replace(' ', '-')})",
                                annotated_content
                            )
                            crossref_map.setdefault(ch_num, []).append({"ref": term, "type": "world"})

                # Update references and content
                latest.references = updated_refs
                latest.content = annotated_content

            # Store cross-reference map in state for UI visualization
            ctx.state.user_feedback["crossref_map"] = json.dumps(crossref_map)

            if inconsistencies:
                ctx.state.error_message = "Cross-reference inconsistencies found:\n" + "\n".join(inconsistencies)
                logger.warning(ctx.state.error_message)
            else:
                logger.info("Cross-reference verification, annotation, and auto-update complete.")
            ctx.state.update_progress("Cross-Reference Verification Complete")
            return GenerateSummaries()
        except Exception as e:
            logger.error(f"Error verifying cross-references: {e}")
            ctx.state.error_message = f"Error verifying cross-references: {e}"
            raise

@dataclass
class RefineWorldWithCharacters(BaseNode[BookGenerationState, BookAgents]):
    """Refines world details based on character profiles."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> RefineCharactersWithWorld:
        freeze = ctx.state.freeze_flags.get("world", False)
        approved = getattr(ctx.state, "approved_characters", True)
        if freeze or not approved:
            logger.info("World refinement frozen or characters not approved, skipping RefineWorldWithCharacters node.")
            ctx.state.update_progress("World Refinement Frozen/Skipped")
            return RefineCharactersWithWorld()

        tasks = BookTasks()
        try:
            refined_world = await run_agent_task(
                ctx,
                "world_builder",
                tasks.build_world,
                ctx.state.book_outline,
                ctx.state.refined_concept,
                update_status="Refining World with Characters"
            )
            ctx.state.world_details = refined_world
            ctx.state.update_progress("World Refined with Characters")
            logger.info("World refined based on character profiles.")
            return RefineCharactersWithWorld()
        except Exception as e:
            logger.error(f"Error refining world: {e}")
            ctx.state.error_message = f"Error refining world: {e}"
            raise

@dataclass
class RefineCharactersWithWorld(BaseNode[BookGenerationState, BookAgents]):
    """Refines character profiles based on world details."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> WriteCoordinator:
        freeze = ctx.state.freeze_flags.get("characters", False)
        approved = getattr(ctx.state, "approved_characters", True)
        if freeze or not approved:
            logger.info("Character refinement frozen or not approved, skipping RefineCharactersWithWorld node.")
            ctx.state.update_progress("Character Refinement Frozen/Skipped")
            return WriteCoordinator()

        tasks = BookTasks()
        try:
            refined_characters_text = await run_agent_task(
                ctx,
                "character_developer",
                tasks.develop_characters,
                ctx.state.book_outline,
                ctx.state.world_details,
                update_status="Refining Characters with World"
            )
            # Parsing logic (reuse from DevelopCharacters)
            parsed_characters = []
            try:
                # Simple example: Assume characters are separated by "---" and fields by lines
                raw_profiles = refined_characters_text.split('---')
                for profile_text in raw_profiles:
                    lines = [line.strip() for line in profile_text.strip().split('\n') if line.strip()]
                    if len(lines) >= 4:
                        parsed_characters.append(CharacterProfile(
                            name=lines[0].replace("Name:", "").strip(),
                            role=lines[1].replace("Role:", "").strip(),
                            description=lines[2].replace("Description:", "").strip(),
                            motivations=lines[3].replace("Motivations:", "").strip(),
                            background="\n".join(lines[4:]) # Crude background capture
                        ))
            except Exception as parse_error:
                logger.error(f"Failed to parse refined character data: {parse_error}")
            ctx.state.characters = parsed_characters
            ctx.state.update_progress("Characters Refined with World")
            logger.info("Characters refined based on world details.")
            return WriteCoordinator()
        except Exception as e:
            logger.error(f"Error refining characters: {e}")
            ctx.state.error_message = f"Error refining characters: {e}"
            raise

# Helper function to run agent tasks within nodes
async def run_agent_task(
    ctx: GraphRunContext[BookGenerationState, BookAgents],
    agent_role: str,
    task_func,
    *args,
    update_status: str = "Running Agent"
) -> str:
    """Runs an agent task with streaming support."""
    # Map agent_role to default system prompts
    default_prompts = {
        "writer": "You are a gifted writer with a command of language that brings stories to life...",
        "reviewer": "You are a professional book reviewer with keen critical analysis skills...",
        "editor": "You are a professional editor who perfects written works...",
        "consistency_checker": "You are a meticulous consistency checker responsible for ensuring narrative coherence...",
        "style_refiner": "You are a sophisticated style refiner dedicated to elevating the quality of prose...",
        "flow_enhancer": "You are an expert in narrative structure and flow, focused on optimizing the reader's journey...",
        "final_formatter": "You are a meticulous formatting specialist responsible for the final presentation of the book...",
        "dialogue_specialist": "You are an expert dialogue writer focused on creating natural, engaging, and character-consistent conversations...",
        "description_architect": "You are a master of descriptive writing, skilled at creating immersive, sensory-rich environments...",
        "plot_continuity_guardian": "You are a meticulous plot continuity expert responsible for maintaining narrative coherence...",
        "market_analyst": "You are an expert in book market analysis. Provide insights on current trends, audience preferences, and competitive landscape.",
        "genre_specialist": "You are a genre expert. Determine the optimal genre positioning for the book based on concept and market data.",
        "audience_analyst": "You analyze reader demographics and preferences to identify the ideal target audience segments.",
        "comparative_titles_researcher": "You research comparable titles in the market to inform positioning and differentiation.",
        "strategic_outliner": "You generate a strategic outline for the book, incorporating market insights and genre conventions.",
        "book_outliner": "You are a skilled story architect with deep understanding of narrative structure...",
        "character_developer": "You are a character development expert who creates memorable, multi-dimensional characters...",
        "world_builder": "You are a master world builder who creates immersive fictional worlds...",
    }
    system_prompt = default_prompts.get(agent_role, f"You are an AI agent specializing in {agent_role} tasks.")

    # --- Adaptive model switching ---
    try:
        # Estimate content length from prompt_data if available
        content_length = 0
        if isinstance(task_details, dict):
            prompt_text = task_details.get("prompt", "")
            content_length = len(prompt_text.split())
        ctx.deps.update_model_map(agent_role, content_length=content_length, complexity=1)
    except Exception:
        pass

    agent = ctx.deps.get_adaptive_agent(agent_role, system_prompt)

    # --- Inject explicit feedback data into agent ---
    try:
        agent.feedback_data = {
            "review_feedback": getattr(ctx.state, "review_feedback", None),
            "feedback_history": [fb.model_dump() for fb in getattr(ctx.state, "feedback_history", [])],
            "user_feedback": getattr(ctx.state, "user_feedback", {}),
            "quality_metrics": ctx.state.user_feedback.get("quality_metrics") if hasattr(ctx.state, "user_feedback") else None,
            "combined_summary": getattr(ctx.state, "combined_summary", None)
        }
    except Exception:
        pass
    start_time = time.time()
    ctx.state.update_progress(f"{update_status}: {agent.role}")
    logger.info(f"Node running agent: {agent.role}")
    
    try:
        if asyncio.iscoroutinefunction(task_func):
            task_details = await task_func(*args)
        else:
            task_details = task_func(*args)
    except Exception as e:
        logger.error(f"Error generating prompt for agent {agent.role}: {e}", exc_info=True)
        ctx.state.error_message = f"Prompt generation failed for agent {agent.role}: {e}"
        ctx.state.update_progress("Error")
        raise

    prompt = task_details.get("prompt", "No prompt generated")

    # --- Inject dynamic context based on feedback and metrics ---
    dynamic_context_parts = []
    try:
        # Prior review feedback
        if hasattr(ctx.state, "review_feedback") and ctx.state.review_feedback:
            dynamic_context_parts.append(f"Prior Review Feedback:\n{ctx.state.review_feedback}")

        # Recent polishing feedback
        if hasattr(ctx.state, "feedback_history") and ctx.state.feedback_history:
            feedback_summaries = "\n".join(
                f"- {fb.category}: {fb.suggestion}" for fb in ctx.state.feedback_history[-5:]
            )
            dynamic_context_parts.append(f"Recent Polishing Feedback:\n{feedback_summaries}")

        # Quality metrics
        if hasattr(ctx.state, "user_feedback") and ctx.state.user_feedback.get("quality_metrics"):
            dynamic_context_parts.append(f"Quality Metrics:\n{ctx.state.user_feedback.get('quality_metrics')}")

        # User feedback
        if hasattr(ctx.state, "user_feedback") and ctx.state.user_feedback:
            user_fb = "\n".join(
                f"{k}: {v}" for k, v in ctx.state.user_feedback.items() if v
            )
            if user_fb:
                dynamic_context_parts.append(f"User Feedback:\n{user_fb}")

        # Combined summary
        if hasattr(ctx.state, "combined_summary") and ctx.state.combined_summary:
            dynamic_context_parts.append(f"Combined Summary:\n{ctx.state.combined_summary}")

    except Exception:
        pass

    dynamic_context = "\n\n".join(dynamic_context_parts)

    # --- Prepend style prompt, then dynamic context, then task prompt ---
    parts = []
    if hasattr(ctx.state, "style_prompt") and ctx.state.style_prompt:
        parts.append(ctx.state.style_prompt)
    if dynamic_context:
        parts.append(dynamic_context)
    parts.append(prompt)
    prompt = "\n\n".join(parts)

    logger.info(f"Prompt for agent '{agent.role}': {prompt}")

    # Emit prompt to callback for full transparency
    if ctx.state.progress_callback:
        try:
            await ctx.state.progress_callback(
                stage=update_status,
                message=f"Prompt sent to {agent.role}",
                agent_name=agent.role,
                percent=0,
                stats={
                    'current_content': f"[Prompt]\n{prompt}",
                    'streaming': True
                }
            )
        except Exception:
            pass

    try:
        # Handle both streaming and non-streaming cases
        if hasattr(agent, 'run_stream'):
            # Initialize streaming buffer
            content_buffer = []
            stream_result = agent.run_stream(prompt)
            
            # Handle different stream result types
            if hasattr(stream_result, 'stream'):
                # Proper async generator case
                async for chunk in stream_result.stream():
                    if isinstance(chunk, str):
                        content_buffer.append(chunk)
                        partial_content = ''.join(content_buffer)

                        # Calculate word count
                        word_count = len(partial_content.split())

                        # Update writing stats
                        now = time.time()
                        elapsed = now - start_time
                        ctx.state.writing_stats.total_writing_time = elapsed
                        ctx.state.writing_stats.total_words_generated = word_count
                        ctx.state.writing_stats.average_wpm = int(word_count / (elapsed / 60)) if elapsed > 0 else 0
                        ctx.state.writing_stats.status = update_status

                        # Maintain cumulative stats during streaming
                        if not hasattr(ctx, "_stream_cumulative"):
                            ctx._stream_cumulative = {
                                'start_time': start_time,
                                'initial_words': ctx.state.total_words_generated,
                                'initial_writing_time': ctx.state.writing_stats.total_writing_time
                            }

                        initial_words = ctx._stream_cumulative['initial_words']
                        initial_writing_time = ctx._stream_cumulative['initial_writing_time']
                        total_words_generated = initial_words + word_count
                        total_writing_time = initial_writing_time + elapsed
                        elapsed_time_total = time.time() - ctx._stream_cumulative['start_time']

                        # --- Update state cumulative stats ---
                        ctx.state.total_words_generated = total_words_generated
                        ctx.state.writing_stats.total_writing_time = total_writing_time
                        ctx.state.writing_stats.average_wpm = (total_words_generated / total_writing_time) * 60 if total_writing_time > 0 else 0

                        try:
                            ctx.state.writing_stats.completed_chapters = sum(
                                1 for status in ctx.state.current_chapter_status.values()
                                if status in ("approved", "draft")
                            )
                        except Exception:
                            ctx.state.writing_stats.completed_chapters = 0

                        try:
                            ctx.state.writing_stats.total_chapters = ctx.state.config.num_chapters or 0
                        except Exception:
                            ctx.state.writing_stats.total_chapters = 0

                        completed_chapters = ctx.state.writing_stats.completed_chapters
                        total_chapters = ctx.state.writing_stats.total_chapters

                        words_per_minute = 0
                        try:
                            words_per_minute = (total_words_generated / elapsed_time_total) * 60 if elapsed_time_total > 0 else 0
                        except Exception:
                            pass

                        average_wpm = ctx.state.writing_stats.average_wpm

                        # Update progress with cumulative stats
                        if ctx.state.progress_callback:
                            try:
                                await ctx.state.progress_callback(
                                    stage=update_status,
                                    message=f"Generating content...",
                                    agent_name=agent.role,
                                    percent=0,  # Can be refined if total target known
                                    stats={
                                        'current_content': partial_content,
                                        'streaming': True,
                                        'word_count': word_count,
                                        'elapsed_time': elapsed_time_total,
                                        'words_generated': total_words_generated,
                                        'words_per_minute': words_per_minute,
                                        'average_wpm': average_wpm,
                                        'completed_chapters': completed_chapters,
                                        'total_chapters': total_chapters,
                                        'total_writing_time': total_writing_time,
                                        'status': update_status
                                    }
                                )
                            except Exception:
                                pass
            else:
                # Fallback to regular run if streaming not properly implemented
                result = await agent.run(prompt)
                content_buffer.append(result.data if hasattr(result, 'data') else str(result))
            
            content = ''.join(content_buffer)
        else:
            # Fallback to regular run if streaming not available
            result = await agent.run(prompt)
            content = result.data if hasattr(result, 'data') else str(result)

        logger.info(f"Response from agent '{agent.role}': {content[:500]}")  # Log first 500 chars

        duration = time.time() - start_time
        ctx.state.record_agent_activity(agent.role, duration, len(content.split()))
        return content

    except Exception as e:
        logger.error(f"Error running agent {agent.role}: {e}", exc_info=True)
        ctx.state.error_message = f"Agent {agent.role} failed: {e}"
        ctx.state.update_progress("Error")
        raise

# === Node Definitions ---

@dataclass
class StartGeneration(BaseNode[BookGenerationState, BookAgents]):
    """Initial node to set up the generation state."""
    initial_concept: str
    config: GenerationConfig

    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> Annotated['DevelopConcept', Edge(label="Develop Concept")]:
        ctx.state.initial_concept = self.initial_concept
        ctx.state.config = self.config
        ctx.state.update_progress("Initialized")
        logger.info(f"Starting book generation run_id: {ctx.state.run_id}")
        print(f"[DEBUG] After StartGeneration: initial_concept={ctx.state.initial_concept!r}, refined_concept={ctx.state.refined_concept!r}")
        return DevelopConcept()

@dataclass
class GenerateTitle(BaseNode[BookGenerationState, BookAgents]):
    """Generates the book title."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> DevelopConcept:
        tasks = BookTasks() # Instantiate tasks helper
        title = await run_agent_task(
            ctx,
            "concept_developer",
            tasks.generate_title, # Pass the method itself
            ctx.state.initial_concept,
            update_status="Generating Title"
        )
        # Defensive fix: ensure title is a string
        title = title or ""
        # Basic title cleaning
        title = re.sub(r'^[#*\s]*(?:Title:)?\s*', '', title).strip()
        title = re.sub(r'[^a-zA-Z0-9 \-\']', '', title)
        title_words = title.split()
        title = ' '.join(title_words[:10])
        if len(title_words) < 2:
            title = f"{title} Novel" if title else "Untitled Story"

        ctx.state.book_title = title
        ctx.state.update_progress("Title Generated")
        logger.info(f"Generated Title: {ctx.state.book_title}")
        return DevelopConcept()

@dataclass
class DevelopConcept(BaseNode[BookGenerationState, BookAgents]):
    """Develops and refines the initial book concept."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> Annotated['PlotScaffolding', Edge(label="Plot Scaffolding")]:
        freeze = ctx.state.freeze_flags.get("concept", False)
        approved = getattr(ctx.state, "approved_concept", True)
        if freeze or not approved:
            logger.info("Concept frozen or not approved, skipping DevelopConcept node.")
            ctx.state.update_progress("Concept Frozen/Skipped")
            return PlotScaffolding()

        tasks = BookTasks()
        # First extract key elements from initial concept
        extracted_elements = await run_agent_task(
            ctx,
            "concept_developer",
            tasks.extract_elements,
            ctx.state.initial_concept,
            update_status="Extracting Elements"
        )
        
        # Compose strategic outline and market data context
        strategic_outline = getattr(ctx.state.publishing_metadata, "strategic_outline", "")
        market_analysis = getattr(ctx.state.publishing_metadata, "market_analysis", "")
        genre_positioning = getattr(ctx.state.publishing_metadata, "genre_positioning", "")
        audience = getattr(ctx.state.publishing_metadata, "audience", "")
        comps = getattr(ctx.state.publishing_metadata, "comparable_titles", "")

        # Then develop enhanced concept using extracted elements and publishing metadata
        refined_concept = await run_agent_task(
            ctx,
            "concept_developer",
            tasks.develop_concept,
            ctx.state.initial_concept,
            extracted_elements,
            ctx.state.publishing_metadata,
            update_status="Developing Concept"
        )
        
        # Fallback: if refined concept is empty or null, use initial concept
        if not refined_concept or not refined_concept.strip():
            logger.warning("Refined concept was empty or null, falling back to initial concept.")
            refined_concept = ctx.state.initial_concept

        # Store both extracted elements and refined concept
        ctx.state.extracted_elements = extracted_elements
        ctx.state.refined_concept = refined_concept
        ctx.state.update_progress("Concept Developed")
        logger.info("Refined book concept.")
        # logger.debug(f"Refined Concept: {refined_concept[:200]}...") # Log snippet
        return PlotScaffolding()

@dataclass
class CreateOutline(BaseNode[BookGenerationState, BookAgents]):
    """Creates the book outline."""
    async def run(
        self, ctx: GraphRunContext[BookGenerationState, BookAgents]
    ) -> Union[
        Annotated[BuildWorld, Edge(label="Build World")],
        Annotated[DevelopCharacters, Edge(label="Develop Characters")]
    ]:
        freeze = ctx.state.freeze_flags.get("outline", False)
        approved = getattr(ctx.state, "approved_outline", True)
        if freeze or not approved:
            logger.info("Outline frozen or not approved, skipping CreateOutline node.")
            ctx.state.update_progress("Outline Frozen/Skipped")
            return BuildWorld()

        tasks = BookTasks()
        outline = await run_agent_task(
            ctx,
            "book_outliner",
            tasks.create_book_outline,
            ctx.state.refined_concept,
            ctx.state.config.num_chapters,
            ctx.state.publishing_metadata,
            update_status="Creating Outline"
        )
        outline = outline or ""
        ctx.state.book_outline = outline
        ctx.state.update_progress("Outline Created")
        logger.info("Created book outline.")
        return BuildWorld()

@dataclass
class BuildWorld(BaseNode[BookGenerationState, BookAgents]):
    """Builds the world details based on the outline and concept."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> DevelopCharacters:
        freeze = ctx.state.freeze_flags.get("world", False)
        approved = getattr(ctx.state, "approved_characters", True)
        if freeze or not approved:
            logger.info("World frozen or characters not approved, skipping BuildWorld node.")
            ctx.state.update_progress("World Frozen/Skipped")
            return DevelopCharacters()

        tasks = BookTasks()
        world_details = await run_agent_task(
            ctx,
            "world_builder",
            tasks.build_world,
            ctx.state.book_outline,
            ctx.state.refined_concept,
            ctx.state.publishing_metadata,
            update_status="Building World"
        )
        world_details = world_details or ""
        ctx.state.world_details = world_details
        ctx.state.update_progress("World Built")
        logger.info("Created world details.")
        # logger.debug(f"World Details: {world_details[:200]}...")
        return DevelopCharacters()

@dataclass
class DevelopCharacters(BaseNode[BookGenerationState, BookAgents]):
    """Develops character profiles based on the outline and world."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> WriteChapter:
        freeze = ctx.state.freeze_flags.get("characters", False)
        approved = getattr(ctx.state, "approved_characters", True)
        if freeze or not approved:
            logger.info("Characters frozen or not approved, skipping DevelopCharacters node.")
            ctx.state.update_progress("Characters Frozen/Skipped")
            return WriteChapter(chapter_num=1)

        tasks = BookTasks()
        character_data_text = await run_agent_task(
            ctx,
            "character_developer",
            tasks.develop_characters,
            ctx.state.book_outline,
            ctx.state.world_details,
            ctx.state.publishing_metadata,
            update_status="Developing Characters"
        )
        character_data_text = character_data_text or ""
        # --- Parsing Logic (Example - needs refinement based on agent output) ---
        # This is a placeholder. Robust parsing is needed here.
        # Ideally, the agent returns structured JSON or uses Pydantic output.
        parsed_characters = []
        try:
            # Simple example: Assume characters are separated by "---" and fields by lines
            raw_profiles = character_data_text.split('---')
            for profile_text in raw_profiles:
                lines = [line.strip() for line in profile_text.strip().split('\n') if line.strip()]
                if len(lines) >= 4: # Basic check
                     # Very basic parsing - needs improvement!
                     parsed_characters.append(CharacterProfile(
                         name=lines[0].replace("Name:", "").strip(),
                         role=lines[1].replace("Role:", "").strip(),
                         description=lines[2].replace("Description:", "").strip(),
                         motivations=lines[3].replace("Motivations:", "").strip(),
                         background="\n".join(lines[4:]) # Crude background capture
                     ))
        except Exception as parse_error:
            logger.error(f"Failed to parse character data: {parse_error}")
            # Decide how to handle - proceed with empty list or raise error?
            # ctx.state.error_message = "Failed to parse character data"
            # raise parse_error

        ctx.state.characters = parsed_characters
        ctx.state.update_progress("Characters Developed")
        logger.info(f"Developed {len(ctx.state.characters)} characters.")
        # Start writing from chapter 1
        return WriteChapter(chapter_num=1)

@dataclass
class WriteChapter(BaseNode[BookGenerationState, BookAgents]):
    """Writes a single chapter with streaming and persistence support."""
    chapter_num: int

    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> Union[WriteChapter, 'ReviewBook']:
        trace.get_current_span().add_event("NodeStart", {"node": "WriteChapter", "chapter": self.chapter_num})

        # Determine current version number
        latest_version_obj = ctx.state.get_latest_chapter_version(self.chapter_num)
        current_version = (latest_version_obj.version + 1) if latest_version_obj else 1

        # Compute file path for this version
        chapter_version_dir = os.path.join(
            'novelForge', 'books', 'chapters', ctx.state.run_id, f'chapter_{self.chapter_num}'
        )
        os.makedirs(chapter_version_dir, exist_ok=True)
        versioned_chapter_path = os.path.join(
            chapter_version_dir, f'v{current_version}.md'
        )

        if ctx.state.has_been_saved:
            logger.warning(f"Book already saved - skipping chapter {self.chapter_num} generation")
            return ReviewBook()

        tasks = BookTasks()
        num_chapters_total = ctx.state.config.num_chapters
        ctx.state.current_chapter_writing = self.chapter_num

        # Prepare context
        prev_chapter_latest = ctx.state.get_latest_chapter_version(self.chapter_num - 1)
        previous_chapter_content = prev_chapter_latest.content if prev_chapter_latest else "First chapter"

        # Compose dynamic context from feedback and metrics
        dynamic_context = compose_dynamic_context(ctx)

        chapter_start_time = time.time()

        # --- Multi-agent initial generation with explicit feedback channels ---
        writer_agent = ctx.deps.writer()
        description_agent = ctx.deps.description_architect()
        dialogue_agent = ctx.deps.dialogue_specialist()

        # Compose prompt for writer agent
        writer_prompt = f"""Chapter {self.chapter_num} Draft

Outline:
{ctx.state.book_outline}

Characters:
{ctx.state.characters}

World Details:
{ctx.state.world_details}

Previous Chapter:
{previous_chapter_content}

Minimum Words: {ctx.state.config.min_words_per_chapter}

Please write the next chapter based on the above context.

At the end of your output, explicitly state the POV character in the format:

POV: Character Name
"""

        trace.get_current_span().add_event("AgentCall", {"agent": "Writer", "phase": "initial"})
        base_narrative, writer_feedback = await writer_agent.process_with_feedback(
            writer_agent.run,
            writer_prompt,
            update_status=f"Writing Chapter {self.chapter_num} (Base Narrative)"
        )

        # --- Extract explicit POV marker ---
        pov_character = None
        pov_provenance = "unknown"
        pov_match = re.search(r"^POV:\s*(.+)$", base_narrative, re.MULTILINE)
        if pov_match:
            pov_character = pov_match.group(1).strip()
            pov_provenance = "explicit"
            # Remove POV line from content
            base_narrative = re.sub(r"^POV:\s*.+$", "", base_narrative, flags=re.MULTILINE).strip()
        else:
            # --- Fallback heuristic ---
            first_part = base_narrative[:max(1000, int(0.2 * len(base_narrative)))]
            counts = {}
            for char in ctx.state.characters:
                count = first_part.count(char.name)
                if count > 0:
                    counts[char.name] = count
            if counts:
                pov_character = max(counts, key=counts.get)
                pov_provenance = "heuristic"

        # Compose prompt for description agent
        description_prompt = f"""Chapter {self.chapter_num} Descriptions

Outline:
{ctx.state.book_outline}

Characters:
{ctx.state.characters}

World Details:
{ctx.state.world_details}

Previous Chapter:
{previous_chapter_content}

Minimum Words: {ctx.state.config.min_words_per_chapter}

Please generate vivid, sensory-rich descriptions for this chapter based on the above context.
"""

        trace.get_current_span().add_event("AgentCall", {"agent": "DescriptionArchitect", "phase": "initial"})
        description_content, description_feedback = await description_agent.process_with_feedback(
            description_agent.run,
            description_prompt,
            update_status=f"Writing Chapter {self.chapter_num} (Descriptions)"
        )

        # Compose prompt for dialogue agent
        dialogue_prompt = f"""Chapter {self.chapter_num} Dialogue

Outline:
{ctx.state.book_outline}

Characters:
{ctx.state.characters}

World Details:
{ctx.state.world_details}

Previous Chapter:
{previous_chapter_content}

Minimum Words: {ctx.state.config.min_words_per_chapter}

Please generate engaging, character-consistent dialogue for this chapter based on the above context.
"""

        trace.get_current_span().add_event("AgentCall", {"agent": "DialogueSpecialist", "phase": "initial"})
        dialogue_content, dialogue_feedback = await dialogue_agent.process_with_feedback(
            dialogue_agent.run,
            dialogue_prompt,
            update_status=f"Writing Chapter {self.chapter_num} (Dialogue)"
        )

        # Save initial outputs
        ctx.state.user_feedback[f"chapter_{self.chapter_num}_base_initial"] = base_narrative
        ctx.state.user_feedback[f"chapter_{self.chapter_num}_description_initial"] = description_content
        ctx.state.user_feedback[f"chapter_{self.chapter_num}_dialogue_initial"] = dialogue_content

        # --- Explicit iterative refinement with feedback passing ---
        refined_base = base_narrative
        refined_description = description_content
        refined_dialogue = dialogue_content

        max_iterations = getattr(ctx.state.config, 'max_iterations', 2)
        all_structured_feedback = []

        prev_base = ""
        prev_desc = ""
        prev_dial = ""

        for iteration in range(1, max_iterations + 1):
            # Cap and prioritize feedback list
            combined_feedback = all_structured_feedback[-6:] + [
                {"category": "transition clarity", "priority": "medium", "suggestion": "Improve scene transitions using sensory anchors."},
                {"category": "character motivation", "priority": "high", "suggestion": "Clarify character desires, fears, and relationships in actions/dialogue."},
                {"category": "emotional resonance", "priority": "high", "suggestion": "Add vulnerability moments and symbolic objects."},
                {"category": "world integration", "priority": "medium", "suggestion": "Embed tactile details and respect world rules."},
                {"category": "thematic depth", "priority": "medium", "suggestion": "Tie themes to character choices and stakes."},
                {"category": "pacing", "priority": "medium", "suggestion": "Follow beat sheet: Goal → Obstacle → Choice → Consequence."}
            ]
            combined_feedback = [
                x.model_dump() if hasattr(x, "model_dump") else x
                for x in combined_feedback
            ]
            combined_feedback = sorted(combined_feedback, key=lambda x: x.get("priority", ""), reverse=True)
            combined_feedback = combined_feedback[:6]

            refined_base, writer_feedback = await writer_agent.process_with_feedback(
                writer_agent.run,
                refined_base,
                feedback_data={
                    "description_feedback": refined_description,
                    "dialogue_feedback": refined_dialogue
                },
                feedback_list=combined_feedback
            )
            all_structured_feedback.extend(writer_feedback)

            refined_description, description_feedback = await description_agent.process_with_feedback(
                description_agent.run,
                refined_description,
                feedback_data={
                    "writer_feedback": refined_base,
                    "dialogue_feedback": refined_dialogue
                },
                feedback_list=combined_feedback
            )
            all_structured_feedback.extend(description_feedback)

            refined_dialogue, dialogue_feedback = await dialogue_agent.process_with_feedback(
                dialogue_agent.run,
                refined_dialogue,
                feedback_data={
                    "writer_feedback": refined_base,
                    "description_feedback": refined_description
                },
                feedback_list=combined_feedback
            )
            all_structured_feedback.extend(dialogue_feedback)

            ctx.state.user_feedback[f"chapter_{self.chapter_num}_base_round_{iteration}"] = refined_base
            ctx.state.user_feedback[f"chapter_{self.chapter_num}_description_round_{iteration}"] = refined_description
            ctx.state.user_feedback[f"chapter_{self.chapter_num}_dialogue_round_{iteration}"] = refined_dialogue

            def diff_ratio(a, b):
                a_tokens = set(a.split())
                b_tokens = set(b.split())
                if not a_tokens and not b_tokens:
                    return 0.0
                return len(a_tokens.symmetric_difference(b_tokens)) / max(len(a_tokens), 1)

            base_change = diff_ratio(prev_base, refined_base)
            desc_change = diff_ratio(prev_desc, refined_description)
            dial_change = diff_ratio(prev_dial, refined_dialogue)

            logger.info(f"Chapter {self.chapter_num} iteration {iteration} content change ratios - Base: {base_change:.3f}, Desc: {desc_change:.3f}, Dial: {dial_change:.3f}")
            logger.info(f"Feedback list size: {len(combined_feedback)}. Context size (tokens): Base: {len(refined_base.split())}, Desc: {len(refined_description.split())}, Dial: {len(refined_dialogue.split())}")
            if len(combined_feedback) > 8:
                logger.warning("Feedback list exceeds recommended size; context window may be at risk.")
            if max(len(refined_base.split()), len(refined_description.split()), len(refined_dialogue.split())) > 2048:
                logger.warning("Agent context size exceeds 2048 tokens; risk of LLM context loss.")

            if base_change < 0.02 and desc_change < 0.02 and dial_change < 0.02:
                logger.info(f"Convergence detected at iteration {iteration}, stopping early.")
                break

            prev_base = refined_base
            prev_desc = refined_description
            prev_dial = refined_dialogue

            logger.info(f"Completed WriteChapter refinement iteration {iteration}")

        if hasattr(ctx.state, "feedback_history") and isinstance(ctx.state.feedback_history, list):
            ctx.state.feedback_history.extend(all_structured_feedback[-6:])
        else:
            ctx.state.feedback_history = all_structured_feedback[-6:]

        # --- Run plot continuity check ---
        _ = await run_agent_task(
            ctx, "plot_continuity_guardian", tasks.write_chapter,
            ctx.state.book_outline, ctx.state.characters, ctx.state.world_details,
            self.chapter_num, previous_chapter_content,
            ctx.state.config.min_words_per_chapter, ctx.state,
            update_status=f"Writing Chapter {self.chapter_num} (Continuity Check)"
        )

        # --- Merge outputs ---
        chapter_content = f"{refined_base}\n\n---\n\n{refined_description}\n\n---\n\n{refined_dialogue}"

        # --- Heuristic content checks before review ---
        flags = []
        try:
            # Check for sensory anchors (look for tactile/sensory words)
            if not re.search(r'\b(touch|feel|rough|smooth|cold|warm|smell|scent|taste|hear|sound|noise|texture|scented|fragrant|bitter|sweet|salty|sour)\b', chapter_content, re.IGNORECASE):
                flags.append("Missing sensory anchor (tactile/sensory detail)")

            # Check for explicit character goals/conflicts
            if not re.search(r'\b(goal|desire|want|need|conflict|struggle|problem|challenge|obstacle|choice|decision|decide|resolve|plan|intend|aim|hope|fear)\b', chapter_content, re.IGNORECASE):
                flags.append("Missing explicit character goal/conflict")

            # Check for emotional moments (vulnerability)
            if not re.search(r'\b(tremble|cry|tear|fear|hesitate|doubt|regret|pain|hurt|wound|scar|weak|vulnerable|shiver|sigh|sob|weep|panic|anxious|nervous|blush|flush)\b', chapter_content, re.IGNORECASE):
                flags.append("Missing emotional/vulnerability moment")
        except Exception:
            pass

        ctx.state.flags[f"chapter_{self.chapter_num}"] = flags

        # --- Iterative review and revision loop ---
        max_revisions = 2
        revision_count = 0
        while revision_count < max_revisions:
            review_feedback = await run_agent_task(
                ctx, "reviewer", tasks.review_book,
                chapter_content,
                update_status=f"Reviewing Chapter {self.chapter_num} (Revision {revision_count + 1})"
            )
            review_feedback = review_feedback or ""

            # If heuristic flags exist, prepend to review feedback to prioritize fixes
            if flags:
                flag_text = "\n".join(f"- {flag}" for flag in flags)
                review_feedback = f"**Auto-Detected Issues:**\n{flag_text}\n\n{review_feedback}"

            if "**Priority:** critical" in review_feedback or flags:
                logger.info(f"Issues found in Chapter {self.chapter_num} revision {revision_count + 1}, revising...")

                revision_prompt = f"""Revise the following chapter draft based on this feedback and detected issues:

FEEDBACK AND ISSUES:
{review_feedback}

CHAPTER DRAFT:
{chapter_content}

Revise the chapter to address all critical issues, improve clarity, style, emotional depth, and coherence, while maintaining narrative flow and character consistency."""

                writer_agent = ctx.deps.writer()
                if hasattr(writer_agent, "process_with_feedback"):
                    revised_content, _ = await writer_agent.process_with_feedback(
                        writer_agent.run,
                        revision_prompt,
                        feedback_source=None,
                        feedback_data={"review_feedback": review_feedback}
                    )
                else:
                    revised_content = await run_agent_task(
                        ctx, "writer",
                        lambda prompt: {"prompt": prompt},
                        revision_prompt,
                        update_status=f"Revising Chapter {self.chapter_num} (Revision {revision_count + 1})"
                    )

                chapter_content = revised_content
                revision_count += 1
            else:
                logger.info(f"No critical issues found in Chapter {self.chapter_num} revision {revision_count + 1}, proceeding.")
                break

        # Save chapter version
        try:
            with open(versioned_chapter_path, 'w', encoding='utf-8') as f:
                f.write(chapter_content)
            logger.info(f"Saved chapter {self.chapter_num} version {current_version} to {versioned_chapter_path}")
        except Exception as e:
            logger.error(f"Failed to save chapter {self.chapter_num} v{current_version}: {e}")
            ctx.state.error_message = f"Failed to save chapter {self.chapter_num} v{current_version}: {e}"
            ctx.state.current_chapter_status[self.chapter_num] = "error"
            raise

        # Summarize chapter content
        try:
            summary_prompt = f"""Summarize the following chapter in 3-5 sentences, capturing key plot points, character developments, and important details:

{chapter_content}

Return only the summary text."""
            summary = await run_agent_task(
                ctx, "reviewer",
                lambda prompt: {"prompt": prompt},
                summary_prompt,
                update_status=f"Summarizing Chapter {self.chapter_num}"
            )
            ctx.state.chapter_summaries[self.chapter_num] = summary or ""
            combined = "\n".join(
                ctx.state.chapter_summaries.get(i, "") for i in range(1, self.chapter_num + 1)
            )
            ctx.state.combined_summary = combined
            logger.info(f"Generated summary for Chapter {self.chapter_num}")
        except Exception as e:
            logger.warning(f"Failed to generate summary for Chapter {self.chapter_num}: {e}")

        # Create chapter result
        word_count = len(chapter_content.split())
        duration = time.time() - chapter_start_time
        wpm = int(word_count / (duration / 60)) if duration > 0 else 0

        pov_character = None
        if ctx.state.characters:
            for char in ctx.state.characters:
                if char.name in chapter_content:
                    pov_character = char.name
                    break

        color_code = None
        if pov_character:
            import hashlib
            hash_digest = hashlib.md5(pov_character.encode()).hexdigest()
            color_code = f"#{hash_digest[:6]}"

        tags = []
        if ctx.state.extracted_elements:
            try:
                import json as json_lib
                elements = ctx.state.extracted_elements
                if isinstance(elements, str):
                    elements = json_lib.loads(elements)
                if isinstance(elements, dict):
                    tags.extend(elements.get("themes", []))
                    tags.extend(elements.get("plot_points", []))
            except Exception:
                pass

        chapter_result = ChapterResult(
            chapter_number=self.chapter_num,
            version=current_version,
            content=chapter_content,
            word_count=word_count,
            start_time=chapter_start_time,
            end_time=time.time(),
            duration=duration,
            words_per_minute=wpm,
            status="draft",
            file_path=versioned_chapter_path,
            revision_notes="Initial draft" if current_version == 1 else f"Rewrite based on feedback (v{current_version-1})",
            pov_character=pov_character,
            color_code=color_code,
            tags=tags
        )

        if not hasattr(ctx.state, "writing_stats") or ctx.state.writing_stats is None:
            logger.error("CRITICAL DEBUG: ctx.state.writing_stats is missing or None before update!")
            raise AttributeError("writing_stats not found on state object where expected in WriteChapter")

        ctx.state.writing_stats.total_writing_time += duration
        ctx.state.writing_stats.average_wpm = int(
            ctx.state.total_words_generated / (ctx.state.writing_stats.total_writing_time / 60)
        ) if ctx.state.writing_stats.total_writing_time > 0 else 0
        ctx.state.writing_stats.completed_chapters = sum(
            1 for status in ctx.state.current_chapter_status.values()
            if status in ("approved", "draft")
        )
        ctx.state.add_chapter_version(chapter_result)

        chapters_needing_rewrite = {
            num for num, status in ctx.state.current_chapter_status.items()
            if status == "needs_rewrite"
        }

        if chapters_needing_rewrite:
            next_rewrite_chapter = min(chapters_needing_rewrite)
            logger.info(f"Next step: Rewriting chapter {next_rewrite_chapter}")
            return WriteChapter(chapter_num=next_rewrite_chapter)
        elif self.chapter_num < num_chapters_total:
            logger.info(f"Next step: Writing chapter {self.chapter_num + 1}")
            return WriteChapter(chapter_num=self.chapter_num + 1)
        else:
            logger.info("All chapters written. Proceeding to ReviewBook.")
            ctx.state.current_chapter_writing = None
            return ReviewBook()


@dataclass
class ReviewBook(BaseNode[BookGenerationState, BookAgents]):
    """Coordinates multi-agent review process and proceeds to aggregation."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'ReviewAggregator':
        ctx.state.update_progress("Reviewing Book")
        logger.info("Starting multi-agent review process...")

        # Apply iteration bounds and feedback/context pruning
        max_iterations = getattr(ctx.state.config, 'max_iterations', 2)
        max_feedback_items = getattr(ctx.state.config, 'max_feedback_items', 6)
        iteration = 0
        feedback_history = []
        while iteration < max_iterations:
            logger.info(f"ReviewBook iteration {iteration+1}/{max_iterations}")
            await asyncio.gather(
                PeerReview().run(ctx),
                EditorialReview().run(ctx),
                ConsistencyCheck().run(ctx),
                StyleRefinement().run(ctx),
                FlowEnhancement().run(ctx)
            )
            # Prune feedback to last max_feedback_items
            for key in list(ctx.state.user_feedback.keys()):
                if isinstance(ctx.state.user_feedback[key], list) and len(ctx.state.user_feedback[key]) > max_feedback_items:
                    ctx.state.user_feedback[key] = ctx.state.user_feedback[key][-max_feedback_items:]
            # Log feedback/context size
            logger.info(f"Feedback keys: {list(ctx.state.user_feedback.keys())}")
            logger.info(f"Feedback/context size: {[len(v) if isinstance(v, list) else 1 for v in ctx.state.user_feedback.values()]}")
            if any(len(v) > max_feedback_items for v in ctx.state.user_feedback.values() if isinstance(v, list)):
                logger.warning("Feedback/context exceeds safe limit!")
            iteration += 1
        logger.info("Multi-agent review process completed.")
        ctx.state.update_progress("Review Complete - Aggregating Feedback")
        return ReviewAggregator()

@dataclass
class PeerReview(BaseNode[BookGenerationState, BookAgents]):
    """Performs peer review simulation."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> ReviewAggregator:
        ctx.state.update_progress("Peer Review")
        tasks = BookTasks()
        try:
            full_draft = ctx.state.get_full_draft()
            peer_review = await run_agent_task(
                ctx, "reviewer", tasks.review_book, full_draft, update_status="Peer Review"
            )
            ctx.state.user_feedback["peer_review"] = peer_review
        except Exception as e:
            ctx.state.user_feedback["peer_review"] = f"Peer review failed: {e}"
            logger.warning(f"Peer review failed: {e}")

@dataclass
class EditorialReview(BaseNode[BookGenerationState, BookAgents]):
    """Performs editorial committee review."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> ReviewAggregator:
        ctx.state.update_progress("Editorial Review")
        tasks = BookTasks()
        try:
            full_draft = ctx.state.get_full_draft()
            editorial_review = await run_agent_task(
                ctx, "reviewer", tasks.review_book, full_draft, update_status="Editorial Review"
            )
            ctx.state.user_feedback["editorial_review"] = editorial_review
        except Exception as e:
            ctx.state.user_feedback["editorial_review"] = f"Editorial review failed: {e}"
            logger.warning(f"Editorial review failed: {e}")
        return ReviewAggregator()

@dataclass
class ConsistencyCheck(BaseNode[BookGenerationState, BookAgents]):
    """Performs consistency checking."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> ReviewAggregator:
        ctx.state.update_progress("Consistency Check")
        tasks = BookTasks()
        try:
            full_draft = ctx.state.get_full_draft()
            consistency_result = await run_agent_task(
                ctx, "consistency_checker", tasks.check_consistency,
                full_draft, ctx.state.characters, ctx.state.world_details,
                ctx.state.config.num_chapters, update_status="Consistency Check"
            )
            ctx.state.user_feedback["consistency"] = consistency_result
        except Exception as e:
            ctx.state.user_feedback["consistency"] = f"Consistency check failed: {e}"
            logger.warning(f"Consistency check failed: {e}")
        return ReviewAggregator()

@dataclass
class StyleRefinement(BaseNode[BookGenerationState, BookAgents]):
    """Performs style refinement."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> ReviewAggregator:
        ctx.state.update_progress("Style Refinement")
        tasks = BookTasks()
        try:
            full_draft = ctx.state.get_full_draft()
            style_result = await run_agent_task(
                ctx, "style_refiner", tasks.refine_style,
                full_draft, ctx.state.style_guidelines,
                ctx.state.config.num_chapters, update_status="Style Refinement"
            )
            ctx.state.user_feedback["style"] = style_result
        except Exception as e:
            ctx.state.user_feedback["style"] = f"Style refinement failed: {e}"
            logger.warning(f"Style refinement failed: {e}")
        return ReviewAggregator()

@dataclass
class FlowEnhancement(BaseNode[BookGenerationState, BookAgents]):
    """Performs flow and pacing enhancement."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> ReviewAggregator:
        ctx.state.update_progress("Flow Enhancement")
        tasks = BookTasks()
        try:
            full_draft = ctx.state.get_full_draft()
            flow_result = await run_agent_task(
                ctx, "flow_enhancer", tasks.enhance_narrative_flow,
                full_draft, ctx.state.book_outline,
                ctx.state.config.num_chapters, update_status="Flow Enhancement"
            )
            ctx.state.user_feedback["flow"] = flow_result
        except Exception as e:
            ctx.state.user_feedback["flow"] = f"Flow enhancement failed: {e}"
            logger.warning(f"Flow enhancement failed: {e}")
        return ReviewAggregator()

@dataclass
class ReviewAggregator(BaseNode[BookGenerationState, BookAgents]):
    """Aggregates review feedback, detects conflicts, and decides on rewrites or escalation."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> Union[
        Annotated['AssembleBook', Edge(label='Review OK')],
        Annotated[WriteChapter, Edge(label='Rewrites Needed')]
    ]:
        ctx.state.update_progress("Aggregating Reviews")
        max_iterations = getattr(ctx.state.config, 'max_iterations', 2)
        max_feedback_items = getattr(ctx.state.config, 'max_feedback_items', 6)
        iteration = 0
        feedback_history = []
        while iteration < max_iterations:
            try:
                feedback_data = {
                    'peer_review': ctx.state.user_feedback.get("peer_review", ""),
                    'editorial_review': ctx.state.user_feedback.get("editorial_review", ""),
                    'consistency': ctx.state.user_feedback.get("consistency", ""),
                    'style': ctx.state.user_feedback.get("style", ""),
                    'flow': ctx.state.user_feedback.get("flow", ""),
                }
                ctx.state.review_feedback = json.dumps(feedback_data)
                # Prune feedback to last max_feedback_items
                for key in feedback_data:
                    if isinstance(ctx.state.user_feedback.get(key), list) and len(ctx.state.user_feedback[key]) > max_feedback_items:
                        ctx.state.user_feedback[key] = ctx.state.user_feedback[key][-max_feedback_items:]
                # Log feedback/context size
                logger.info(f"ReviewAggregator feedback/context size: {[len(v) if isinstance(v, list) else 1 for v in ctx.state.user_feedback.values()]}")
                if any(len(v) > max_feedback_items for v in ctx.state.user_feedback.values() if isinstance(v, list)):
                    logger.warning("ReviewAggregator feedback/context exceeds safe limit!")
                # --- Conflict detection ---
                conflict_detected = False
                conflict_reasons = []
                style_feedback = feedback_data.get('style', '').lower()
                flow_feedback = feedback_data.get('flow', '').lower()
                if "more poetic" in style_feedback and "clearer" in flow_feedback:
                    conflict_detected = True
                    conflict_reasons.append("Style wants more poetic, flow wants clearer")
                if "simplify" in style_feedback and "add detail" in flow_feedback:
                    conflict_detected = True
                    conflict_reasons.append("Style wants simplification, flow wants more detail")
                if conflict_detected:
                    logger.info(f"Conflict detected in feedback: {conflict_reasons}")
                    ctx.state.user_feedback["conflict_reasons"] = conflict_reasons
                    editor_agent = ctx.deps.editor()
                    content = ctx.state.get_full_draft()
                    combined_feedback = "\n\n".join(f"{k}: {v}" for k, v in feedback_data.items())
                    arbitration_prompt = f"""The following feedback contains conflicting suggestions:\n\n{conflict_reasons}\n\nFeedback:\n{combined_feedback}\n\nPlease review the draft and produce a balanced, final revision that resolves these conflicts, balancing poetic style with clarity, detail, and flow."""

                    revised_content, _ = await editor_agent.process_with_feedback(
                        editor_agent.run,
                        arbitration_prompt
                    )
                    ctx.state.final_book_content = revised_content
                    ctx.state.update_progress("Editor Arbitration Complete")
                    return AssembleBook()
                # --- Rewrite detection ---
                chapters_to_rewrite = set()
                for category, feedback_text in feedback_data.items():
                    if "**Priority:** critical" in feedback_text:
                        matches = re.findall(r'- \*\*Chapter:\*\* (\d+)', feedback_text)
                        for m in matches:
                            try:
                                ch_num = int(m)
                                if 1 <= ch_num <= ctx.state.config.num_chapters:
                                    chapters_to_rewrite.add(ch_num)
                            except:
                                continue

                for chapter_num in chapters_to_rewrite:
                    if chapter := ctx.state.get_chapter(chapter_num):
                        chapter.status = "needs_rewrite"

                if chapters_to_rewrite:
                    next_rewrite = min(chapters_to_rewrite)
                    ctx.state.update_progress(f"Review Complete - Rewrite Ch {next_rewrite}")
                    return WriteChapter(chapter_num=next_rewrite)
                else:
                    approved_count = 0
                    for ch_num in range(1, ctx.state.config.num_chapters + 1):
                        ctx.state.current_chapter_status[ch_num] = "approved"
                        latest_chapter = ctx.state.get_latest_chapter_version(ch_num)
                        if latest_chapter:
                            latest_chapter.status = "approved"
                            approved_count += 1
                    ctx.state.writing_stats.chapters_approved = approved_count
                    ctx.state.update_progress("Review Complete - All Approved")
                    return AssembleBook()
            except Exception as e:
                ctx.state.error_message = f"Review aggregation failed: {e}"
                logger.warning(f"Review aggregation failed: {e}")
                if iteration == max_iterations - 1:
                    return AssembleBook()
            iteration += 1
        return AssembleBook()

@dataclass
class AssembleBook(BaseNode[BookGenerationState, BookAgents]):
    """Assembles approved chapters and metadata into a single manuscript string."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'PolishBook':
        logger.info("Assembling approved chapters and metadata into manuscript...")
        ctx.state.update_progress("Assembling Book")

        approved_chapters = []
        for i in range(1, ctx.state.config.num_chapters + 1):
            ch = ctx.state.get_latest_approved_chapter(i)
            if ch:
                approved_chapters.append(ch)
            else:
                logger.warning(f"Chapter {i} missing approved version, excluding from assembly.")

        if not approved_chapters:
            ctx.state.error_message = "No approved chapters found during assembly"
            ctx.state.update_progress("Error")
            return PolishBook()

        toc = "\n".join(
            f"{idx + 1}. Chapter {chap.chapter_number}"
            for idx, chap in enumerate(approved_chapters)
        )

        manuscript = f"# {ctx.state.book_title}\n\n"
        manuscript += f"## Table of Contents\n{toc}\n\n"
        manuscript += "\n\n".join(
            f"# Chapter {chap.chapter_number}\n\n{chap.content}"
            for chap in approved_chapters
        )

        if ctx.state.refined_concept:
            manuscript += f"\n\n## Key Elements\n\n**Concept:** {ctx.state.refined_concept}\n\n"
        if ctx.state.world_details:
            manuscript += f"**World Details:**\n{ctx.state.world_details}\n\n"
        if ctx.state.characters:
            manuscript += "**Characters:**\n" + "\n".join(
                f"- {char.name}: {char.role}" for char in ctx.state.characters
            ) + "\n"

        ctx.state.final_book_content = manuscript
        ctx.state.update_progress("Book Assembled")
        return PolishBook()

@dataclass
class PolishBook(BaseNode[BookGenerationState, BookAgents]):
    """Polishes the assembled manuscript using specialized agents."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> FormatBook:
        logger.info("Polishing assembled manuscript...")
        ctx.state.update_progress("Polishing Book")
        content = ctx.state.final_book_content or ctx.state.get_full_draft()
        prompt_data = TaskPrompts.polish_book(content)
        try:
            polished = await run_agent_task(
                ctx, "book_polisher", lambda *_: prompt_data,
                update_status="Polishing Book"
            )
            ctx.state.final_book_content = polished
        except Exception as e:
            logger.error(f"Book polishing failed: {e}")
        ctx.state.update_progress("Polishing Complete")
        return FormatBook()

@dataclass
class FormatBook(BaseNode[BookGenerationState, BookAgents]):
    """Formats the polished manuscript for final presentation."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> 'SaveFinalBook':
        logger.info("Formatting polished manuscript...")
        ctx.state.update_progress("Formatting Book")
        content = ctx.state.final_book_content or ""
        prompt_data = TaskPrompts.format_book(content)
        try:
            formatted = await run_agent_task(
                ctx, "final_formatter", lambda *_: prompt_data,
                update_status="Formatting Final Book"
            )
            ctx.state.final_book_content = formatted
        except Exception as e:
            logger.error(f"Formatting failed: {e}")
        ctx.state.update_progress("Formatting Complete")
        return SaveFinalBook()

@dataclass
class SaveFinalBook(BaseNode[BookGenerationState, BookAgents]):
    """Saves the final polished and formatted manuscript."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> ExportEPUB:
        logger.info("Saving final polished and formatted book...")
        ctx.state.update_progress("Saving Final Book")

        # Propagate title from metadata if missing
        if not ctx.state.book_title and getattr(ctx.state.publishing_metadata, "title", None):
            ctx.state.book_title = ctx.state.publishing_metadata.title

        if ctx.state.has_been_saved:
            logger.warning("Book already saved, skipping save.")
            return ExportEPUB()

        safe_title = re.sub(r'[^\w\-_\. ]', '_', ctx.state.book_title or "Untitled")[:50]
        timestamp = time.strftime('%Y%m%d_%H%M%S')
        final_dir = os.path.join('novelForge', 'books', 'drafts')
        os.makedirs(final_dir, exist_ok=True)

        final_path = os.path.join(final_dir, f'{safe_title}_{timestamp}.md')
        stats_path = os.path.join(final_dir, f'{safe_title}_{timestamp}_stats.json')

        try:
            with open(final_path, 'w', encoding='utf-8') as f:
                f.write(ctx.state.final_book_content or "")

            approved_chapters = [
                ctx.state.get_latest_approved_chapter(i)
                for i in range(1, ctx.state.config.num_chapters + 1)
                if ctx.state.get_latest_approved_chapter(i)
            ]
            approved_word_count = sum(ch.word_count for ch in approved_chapters)

            stats = {
                'title': ctx.state.book_title,
                'chapters_requested': ctx.state.config.num_chapters,
                'chapters_approved_included': len(approved_chapters),
                'approved_chapter_numbers': [ch.chapter_number for ch in approved_chapters],
                'word_counts_approved': [ch.word_count for ch in approved_chapters],
                'total_words_approved': approved_word_count,
                'generation_time': time.time() - ctx.state.start_time,
                'final_file_path': final_path,
                'approved_chapter_files': [ch.file_path for ch in approved_chapters if ch.file_path]
            }
            with open(stats_path, 'w', encoding='utf-8') as f:
                json.dump(stats, f, indent=2)

            ctx.state.final_book_path = final_path
            ctx.state.final_stats_path = stats_path
            ctx.state.has_been_saved = True
            ctx.state.end_time = time.time()
            ctx.state.update_progress("Book Saved")
            logger.info(f"Final book saved to: {final_path}")
        except Exception as e:
            logger.error(f"Failed to save final book: {e}")
            ctx.state.error_message = f"Failed to save final book: {e}"
            ctx.state.update_progress("Error")
        return ExportEPUB()

@dataclass
class ExportEPUB(BaseNode[BookGenerationState, BookAgents]):
    """Exports the saved Markdown book to EPUB format."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> ExportPDF:
        logger.info("Exporting book to EPUB...")
        ctx.state.update_progress("Exporting EPUB")
        content = ctx.state.final_book_content or ctx.state.get_full_draft()
        prompt_data = TaskPrompts.export_epub(content, ctx.state.book_title, getattr(ctx.state.publishing_metadata, "author", None))
        try:
            epub_path = await run_agent_task(
                ctx, "epub_exporter", lambda *_: prompt_data,
                update_status="Exporting EPUB"
            )
            ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"epub_path": epub_path, "epub_export_status": "success"})
        except Exception as e:
            ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"epub_path": None, "epub_export_status": "failed"})
            logger.error(f"EPUB export failed: {e}")
        return ExportPDF()

@dataclass
class ExportPDF(BaseNode[BookGenerationState, BookAgents]):
    """Exports the saved Markdown book to PDF format."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> ExportDOCX:
        logger.info("Exporting book to PDF...")
        ctx.state.update_progress("Exporting PDF")
        content = ctx.state.final_book_content or ctx.state.get_full_draft()
        try:
            import markdown2
            from fpdf import FPDF
            import os

            md_path = ctx.state.final_book_path
            if not md_path or not os.path.exists(md_path):
                logger.warning("Markdown file not found for PDF export.")
                ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"pdf_path": None})
                ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"pdf_export_status": "failed"})
                return ExportDOCX()

            with open(md_path, 'r', encoding='utf-8') as f:
                md_text = f.read()

            html = markdown2.markdown(md_text, extras=["fenced-code-blocks", "tables", "header-ids"])

            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_title(ctx.state.book_title or "Untitled")
            pdf.set_author(getattr(ctx.state.publishing_metadata, "author", None) or "AI Author")

            # Embed Unicode font for full UTF-8 support
            try:
                font_path = os.path.join("assets", "DejaVuSans.ttf")
                pdf.add_font("DejaVu", "", font_path, uni=True)
                pdf.set_font("DejaVu", "", 12)
            except Exception:
                # Fallback to Arial if font not found
                pdf.set_font("Arial", size=12)

            # Simple HTML parsing: split by lines, add headings and paragraphs
            for line in html.splitlines():
                line = line.strip()
                if not line:
                    pdf.ln(5)
                    continue
                if line.startswith("<h1"):
                    pdf.set_font("Arial", 'B', 16)
                    text = re.sub('<[^<]+?>', '', line)
                    pdf.cell(0, 10, text, ln=True)
                    pdf.set_font("Arial", size=12)
                elif line.startswith("<h2"):
                    pdf.set_font("Arial", 'B', 14)
                    text = re.sub('<[^<]+?>', '', line)
                    pdf.cell(0, 10, text, ln=True)
                    pdf.set_font("Arial", size=12)
                elif line.startswith("<h3"):
                    pdf.set_font("Arial", 'B', 12)
                    text = re.sub('<[^<]+?>', '', line)
                    pdf.cell(0, 10, text, ln=True)
                    pdf.set_font("Arial", size=12)
                else:
                    text = re.sub('<[^<]+?>', '', line)
                    pdf.multi_cell(0, 10, text)

            pdf_filename = os.path.basename(md_path).replace('.md', '.pdf')
            pdf_dir = os.path.join('novelForge', 'books', 'drafts')
            os.makedirs(pdf_dir, exist_ok=True)
            pdf_path = os.path.join(pdf_dir, pdf_filename)
            pdf.output(pdf_path)
            ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"pdf_path": os.path.abspath(pdf_path)})
            ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"pdf_export_status": "success"})
            logger.info(f"Exported PDF to {pdf_path}")
        except Exception as e:
            logger.warning(f"PDF export failed: {e}")
            try:
                ctx.state.publishing_metadata.pdf_path = None
                ctx.state.publishing_metadata.pdf_export_status = 'failed'
            except Exception:
                pass
        return ExportDOCX()

@dataclass
class ExportDOCX(BaseNode[BookGenerationState, BookAgents]):
    """Exports the saved Markdown book to DOCX format."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> FinalReport:
        ctx.state.update_progress("Exporting DOCX")
        try:
            import markdown2
            import docx
            import os

            md_path = ctx.state.final_book_path
            if not md_path or not os.path.exists(md_path):
                logger.warning("Markdown file not found for DOCX export.")
                ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"docx_path": None})
                ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"docx_export_status": "failed"})
                return FinalReport()

            with open(md_path, 'r', encoding='utf-8') as f:
                md_text = f.read()

            html = markdown2.markdown(md_text, extras=["fenced-code-blocks", "tables", "header-ids"])

            doc = docx.Document()

            # Set core properties
            core_props = doc.core_properties
            core_props.title = ctx.state.book_title or "Untitled"
            core_props.author = getattr(ctx.state.publishing_metadata, "author", None) or "AI Author"

            # Parse HTML lines and add to DOCX
            for line in html.splitlines():
                line = line.strip()
                if not line:
                    continue
                if line.startswith("<h1"):
                    doc.add_heading(text=re.sub('<[^<]+?>', '', line), level=1)
                elif line.startswith("<h2"):
                    doc.add_heading(text=re.sub('<[^<]+?>', '', line), level=2)
                elif line.startswith("<h3"):
                    doc.add_heading(text=re.sub('<[^<]+?>', '', line), level=3)
                else:
                    text = re.sub('<[^<]+?>', '', line)
                    doc.add_paragraph(text)

            docx_filename = os.path.basename(md_path).replace('.md', '.docx')
            docx_dir = os.path.join('novelForge', 'books', 'drafts')
            os.makedirs(docx_dir, exist_ok=True)
            docx_path = os.path.join(docx_dir, docx_filename)
            doc.save(docx_path)
            ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"docx_path": os.path.abspath(docx_path)})
            ctx.state.publishing_metadata = ctx.state.publishing_metadata.model_copy(update={"docx_export_status": "success"})
            logger.info(f"Exported DOCX to {docx_path}")
        except Exception as e:
            logger.warning(f"DOCX export failed: {e}")
            try:
                ctx.state.publishing_metadata.docx_path = None
                ctx.state.publishing_metadata.docx_export_status = 'failed'
            except Exception:
                pass
        return FinalReport()

@dataclass
class FinalReport(BaseNode[BookGenerationState, BookAgents]):
    """Generates a final summary report of the book generation process."""
    async def run(self, ctx: GraphRunContext[BookGenerationState, BookAgents]) -> End[str]:
        ctx.state.update_progress("Generating Final Report")
        try:
            report = f"# Novel Forge Generation Report\n\n"
            report += f"## Title\n{ctx.state.book_title or 'Untitled'}\n\n"
            report += f"## Concept\n{ctx.state.refined_concept or ctx.state.initial_concept}\n\n"
            report += f"## Outline\n{ctx.state.book_outline or 'N/A'}\n\n"
            report += f"## Market Insights\n"
            report += f"- Market Analysis: {getattr(ctx.state.publishing_metadata, 'market_analysis', 'N/A')}\n"
            report += f"- Genre Positioning: {getattr(ctx.state.publishing_metadata, 'genre_positioning', 'N/A')}\n"
            report += f"- Audience: {getattr(ctx.state.publishing_metadata, 'audience', 'N/A')}\n"
            report += f"- Comparable Titles: {getattr(ctx.state.publishing_metadata, 'comparable_titles', 'N/A')}\n"
            report += f"- Proposal: {getattr(ctx.state.publishing_metadata, 'proposal', 'N/A')}\n\n"
            report += f"## Exported Files\n"
            report += f"- Markdown: {ctx.state.final_book_path or 'N/A'}\n"
            report += f"- EPUB: {getattr(ctx.state.publishing_metadata, 'epub_path', 'N/A')}\n"
            report += f"- PDF: {getattr(ctx.state.publishing_metadata, 'pdf_path', 'N/A')}\n"
            report += f"- DOCX: {getattr(ctx.state.publishing_metadata, 'docx_path', 'N/A')}\n\n"
            report += f"## Metrics\n"
            report += f"- Total Words: {ctx.state.total_words_generated}\n"
            report += f"- Chapters: {ctx.state.config.num_chapters}\n"
            report += f"- Generation Time (s): {time.time() - ctx.state.start_time:.2f}\n\n"
            if ctx.state.error_message:
                report += f"## Errors/Warnings\n{ctx.state.error_message}\n\n"
            else:
                report += "## Errors/Warnings\nNone\n\n"
            ctx.state.publishing_metadata['final_report'] = report
            logger.info("Final report generated.")
        except Exception as e:
            logger.warning(f"Failed to generate final report: {e}")
            ctx.state.publishing_metadata['final_report'] = f"Failed to generate report: {e}"
        return End(ctx.state.final_book_path or "Generation complete")

def compose_dynamic_context(ctx):
    """
    Compose a robust, agentic dynamic context for chapter writing and downstream nodes.
    Includes feedback, metrics, prior chapter summaries, and any relevant state fields.
    Strictly follows pydantic-ai agent framework best practices.
    """
    state = ctx.state
    # Gather prior feedback
    feedback = getattr(state, "feedback", "") or "None"
    # Gather advanced quality metrics if available
    metrics = getattr(state, "quality_metrics", {})
    metrics_str = "\n".join(f"{k}: {v}" for k, v in metrics.items()) if metrics else "None"
    # Gather prior summaries
    summaries = getattr(state, "combined_summary", "") or "None"
    # Gather agent review notes
    agent_notes = getattr(state, "agent_notes", "") or "None"
    # Gather style guidelines
    style_guidelines = getattr(state, "style_guidelines", "") or "None"
    # Compose context string
    context = f"""
--- DYNAMIC CONTEXT ---
Feedback: {feedback}
Quality Metrics:\n{metrics_str}
Prior Summaries:\n{summaries}
Agent Notes: {agent_notes}
Style Guidelines: {style_guidelines}
----------------------
"""
    return context
